#!/usr/bin/env python3
"""Validate browser-downloaded conversion artifacts."""

from __future__ import annotations

import argparse
import json
import re
import zipfile
from pathlib import Path

import fitz
from docx import Document
from docx.oxml.ns import qn


def image_non_white_ratio(content: bytes) -> float:
    pixmap = fitz.Pixmap(content)
    samples = pixmap.samples
    non_white = 0
    for index in range(0, len(samples), pixmap.n):
        opaque = pixmap.n < 4 or samples[index + pixmap.n - 1] >= 32
        if opaque and any(
            samples[index + channel] < 248
            for channel in range(min(3, pixmap.n))
        ):
            non_white += 1
    return round(non_white / (pixmap.width * pixmap.height), 5)


def validate_docx(path: Path) -> dict[str, object]:
    if not zipfile.is_zipfile(path):
        raise ValueError("Le DOCX n'est pas une archive ZIP valide.")
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    media_parts = [
        part
        for part in document.part.package.parts
        if part.content_type.startswith("image/")
        and "thumbnail" not in str(part.partname)
    ]
    image_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph._p.xpath(".//w:drawing")
    ]
    exact_line_rule_paragraphs = 0
    for paragraph in image_paragraphs:
        spacing_nodes = paragraph._p.xpath("./w:pPr/w:spacing")
        if spacing_nodes and spacing_nodes[0].get(qn("w:lineRule")) == "exact":
            exact_line_rule_paragraphs += 1
    image_extents = []
    for index, shape in enumerate(document.inline_shapes):
        section = document.sections[min(index, len(document.sections) - 1)]
        image_extents.append(
            {
                "widthPt": round(shape.width.pt, 2),
                "heightPt": round(shape.height.pt, 2),
                "pageWidthRatio": round(
                    shape.width / section.page_width,
                    4,
                ),
                "pageHeightRatio": round(
                    shape.height / section.page_height,
                    4,
                ),
            }
        )
    return {
        "valid": True,
        "text": text,
        "paragraphCount": len(document.paragraphs),
        "imageCount": len(document.inline_shapes),
        "tableCount": len(document.tables),
        "sectionCount": len(document.sections),
        "imageNonWhiteRatios": [
            image_non_white_ratio(part.blob) for part in media_parts
        ],
        "imageExtents": image_extents,
        "imageParagraphCount": len(image_paragraphs),
        "exactLineRuleImageParagraphs": exact_line_rule_paragraphs,
        "clippingDetected": exact_line_rule_paragraphs > 0,
        "orientations": [
            "landscape"
            if section.page_width > section.page_height
            else "portrait"
            for section in document.sections
        ],
    }


def validate_txt(path: Path) -> dict[str, object]:
    text = path.read_bytes().decode("utf-8")
    return {
        "valid": True,
        "text": text,
        "pageSeparators": text.count("--- Page "),
    }


def validate_html(path: Path) -> dict[str, object]:
    text = path.read_bytes().decode("utf-8")
    external_resources = re.findall(
        r"""(?:src|href)=["']https?://""",
        text,
        flags=re.IGNORECASE,
    )
    return {
        "valid": text.lower().startswith("<!doctype html>"),
        "text": re.sub("<[^>]+>", " ", text),
        "pageSections": text.count('class="pdf-page"'),
        "externalResourceCount": len(external_resources),
        "embeddedImageCount": text.count("data:image/"),
    }


def image_details(content: bytes) -> dict[str, object]:
    pixmap = fitz.Pixmap(content)
    return {
        "width": pixmap.width,
        "height": pixmap.height,
        "colorspace": pixmap.colorspace.name if pixmap.colorspace else None,
    }


def validate_images(path: Path) -> dict[str, object]:
    if zipfile.is_zipfile(path):
        with zipfile.ZipFile(path) as archive:
            names = archive.namelist()
            images = [image_details(archive.read(name)) for name in names]
        return {
            "valid": True,
            "archive": True,
            "names": names,
            "images": images,
        }
    return {
        "valid": True,
        "archive": False,
        "names": [path.name],
        "images": [image_details(path.read_bytes())],
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("artifact", type=Path)
    parser.add_argument(
        "--format",
        required=True,
        choices=["docx", "txt", "html", "png", "jpeg"],
    )
    args = parser.parse_args()

    if args.format == "docx":
        result = validate_docx(args.artifact)
    elif args.format == "txt":
        result = validate_txt(args.artifact)
    elif args.format == "html":
        result = validate_html(args.artifact)
    else:
        result = validate_images(args.artifact)
    result["size"] = args.artifact.stat().st_size
    print(json.dumps(result))


if __name__ == "__main__":
    main()
