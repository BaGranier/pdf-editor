from __future__ import annotations

import hashlib
import json
import os
import shutil
import subprocess
from pathlib import Path
from zipfile import ZipFile

import fitz
import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.conversion.docx_converter import (
    PdfToDocxConverter,
    text_retention_ratio,
)


FIXTURE = (
    Path(__file__).parent
    / "fixtures"
    / "docx_regression"
    / "corteva_report"
    / "input.pdf"
)
FIXTURE_SHA256 = "890d95bcb4767569749ac0d7f84ae285ff921208346ef80579b2454ac984b230"
LIBREOFFICE_AVAILABLE = shutil.which("libreoffice") is not None


def _story_field_instructions(story: object) -> list[str]:
    return [
        (element.text or "").strip()
        for element in story._element.xpath(".//w:instrText")
    ]


def _story_is_non_empty(story: object) -> bool:
    return bool(
        "".join(paragraph.text for paragraph in story.paragraphs).strip()
        or story._element.xpath(".//w:drawing | .//w:instrText")
    )


def _body_section_drawing_counts(document: Document) -> list[int]:
    counts: list[int] = []
    current = 0
    for element in document.element.body.iterchildren():
        current += len(element.xpath(".//w:drawing"))
        if element.tag == qn("w:p") and element.xpath("./w:pPr/w:sectPr"):
            counts.append(current)
            current = 0
    counts.append(current)
    return counts


def _render_with_libreoffice(docx_path: Path, output_directory: Path) -> Path:
    profile = output_directory / "lo-profile"
    cache = output_directory / "lo-cache"
    config = output_directory / "lo-config"
    runtime = output_directory / "lo-runtime"
    for directory in (cache, config, runtime):
        directory.mkdir(mode=0o700)
    result = subprocess.run(
        [
            shutil.which("libreoffice") or "libreoffice",
            "--headless",
            f"-env:UserInstallation={profile.as_uri()}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_directory),
            str(docx_path),
        ],
        check=False,
        capture_output=True,
        text=True,
        timeout=90,
        env={
            **os.environ,
            "XDG_CACHE_HOME": str(cache),
            "XDG_CONFIG_HOME": str(config),
            "XDG_RUNTIME_DIR": str(runtime),
        },
    )
    rendered = output_directory / f"{docx_path.stem}.pdf"
    assert result.returncode == 0 and rendered.is_file(), (
        result.stdout,
        result.stderr,
    )
    return rendered


@pytest.mark.integration
@pytest.mark.regression
@pytest.mark.skipif(
    not LIBREOFFICE_AVAILABLE,
    reason="LibreOffice est requis pour la régression Corteva.",
)
def test_corteva_real_pdf_keeps_flow_stories_fields_images_and_three_pages(
    tmp_path: Path,
) -> None:
    """Combined HF/PARA/FLOW regression through the production converter."""

    source_before = FIXTURE.read_bytes()
    assert hashlib.sha256(source_before).hexdigest() == FIXTURE_SHA256
    output = tmp_path / "corteva.docx"
    converter = PdfToDocxConverter()
    converter.convert(FIXTURE, output)
    document = Document(output)
    rendered_path = _render_with_libreoffice(output, tmp_path)

    with fitz.open(FIXTURE) as source, fitz.open(rendered_path) as rendered:
        source_page_count = len(source)
        rendered_page_count = len(rendered)
        source_text = "\n".join(
            " ".join(str(word[4]) for word in page.get_text("words", sort=True))
            for page in source
        )
        rendered_page_image_counts = [
            len(page.get_images(full=True)) for page in rendered
        ]
        rendered_top_image_counts = [
            sum(
                block.get("type") == 1
                and block.get("bbox", (0, page.rect.height))[1]
                < page.rect.height * 0.15
                for block in page.get_text("dict", sort=True).get("blocks", [])
            )
            for page in rendered
        ]
        page_two_footer_blocks = [
            block
            for block in rendered[1].get_text("blocks", sort=True)
            if str(block[4]).strip() == "2"
            and block[1] >= rendered[1].rect.height * 0.85
        ]

    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    story_text = "\n".join(
        element.text or ""
        for section in document.sections
        for story in (section.header, section.footer)
        for element in story._element.xpath(".//w:t")
    )
    first_body_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if "Suite à cette enquête" in paragraph.text
    ]
    list_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.style.name == "List Bullet"
    ]
    header_story_count = len(document.sections)
    non_empty_header_count = sum(
        _story_is_non_empty(section.header) for section in document.sections
    )
    footer_story_count = len(document.sections)
    non_empty_footer_count = sum(
        _story_is_non_empty(section.footer) for section in document.sections
    )
    page_field_count = sum(
        _story_field_instructions(section.footer).count("PAGE")
        for section in document.sections
    )
    numpages_field_count = sum(
        _story_field_instructions(section.footer).count("NUMPAGES")
        for section in document.sections
    )
    meaningful_paragraph_count = sum(
        bool(paragraph.text.strip()) for paragraph in document.paragraphs
    )

    with ZipFile(output) as archive:
        names = set(archive.namelist())
        header_parts = sorted(
            name
            for name in names
            if name.startswith("word/header") and name.endswith(".xml")
        )
        footer_parts = sorted(
            name
            for name in names
            if name.startswith("word/footer") and name.endswith(".xml")
        )
        header_relationship_parts = sorted(
            name
            for name in names
            if name.startswith("word/_rels/header")
            and name.endswith(".xml.rels")
        )
        document_xml = archive.read("word/document.xml").decode("utf-8")
        relationships_xml = archive.read(
            "word/_rels/document.xml.rels"
        ).decode("utf-8")
        header_xml = [archive.read(name).decode("utf-8") for name in header_parts]
        footer_xml = [archive.read(name).decode("utf-8") for name in footer_parts]
        header_relationship_xml = [
            archive.read(name).decode("utf-8")
            for name in header_relationship_parts
        ]

    metrics = {
        "source_page_count": source_page_count,
        "rendered_page_count": rendered_page_count,
        "page_count_delta": rendered_page_count - source_page_count,
        "header_story_count": header_story_count,
        "non_empty_header_count": non_empty_header_count,
        "footer_story_count": footer_story_count,
        "non_empty_footer_count": non_empty_footer_count,
        "PAGE_field_count": page_field_count,
        "NUMPAGES_field_count": numpages_field_count,
        "logical_line_count": converter.last_flow_metrics["logical_lines"],
        "paragraph_count": meaningful_paragraph_count,
        "text_retention_ratio": round(
            text_retention_ratio(source_text, f"{body_text}\n{story_text}"),
            4,
        ),
    }
    print(json.dumps(metrics, ensure_ascii=False, sort_keys=True))

    assert source_page_count == rendered_page_count == len(document.sections) == 3
    assert metrics["page_count_delta"] == 0
    assert len(header_parts) == len(footer_parts) == 3
    assert len(header_relationship_parts) == 3
    assert document_xml.count("headerReference") == 3
    assert document_xml.count("footerReference") == 3
    assert relationships_xml.count("/header") == 3
    assert relationships_xml.count("/footer") == 3
    assert all("<w:drawing" in xml for xml in header_xml)
    assert all("/image" in xml for xml in header_relationship_xml)
    assert all("<w:drawing" not in xml for xml in footer_xml)
    assert non_empty_header_count == 3
    assert non_empty_footer_count == 2
    assert page_field_count == 2
    assert numpages_field_count == 0
    assert all(
        _story_field_instructions(section.footer) == ["PAGE"]
        for section in document.sections[1:]
    )
    assert not _story_field_instructions(document.sections[0].footer)
    assert [
        section._sectPr.xpath("./w:pgNumType")[0].get(qn("w:start"))
        for section in document.sections[1:]
    ] == ["2", "13"]
    assert len({section.header.part.partname for section in document.sections}) == 3
    assert len({section.footer.part.partname for section in document.sections}) == 3
    assert not any(
        line.strip() in {"2", "13"} for line in body_text.splitlines()
    )

    assert len(first_body_paragraphs) == 1
    first_body = first_body_paragraphs[0]
    assert "la prestation de service en application de produits phytosanitaires. Les" in (
        first_body.text
    )
    assert first_body.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert first_body.paragraph_format.first_line_indent is not None
    assert first_body.paragraph_format.first_line_indent.pt == pytest.approx(36)
    assert any(run.italic for run in first_body.runs if "répondants" in run.text)
    assert len(list_paragraphs) == 4
    assert [paragraph.text for paragraph in list_paragraphs] == [
        "le Grand-Est",
        "les Hauts de France / Ile de France",
        "la Bretagne / Pays de Loire",
        "l’Occitanie / Aquitaine",
    ]
    assert converter.last_flow_metrics["same_baseline_fragments_merged"] >= 13
    assert _body_section_drawing_counts(document) == [0, 1, 1]

    assert rendered_top_image_counts == [1, 1, 1]
    assert rendered_page_image_counts[0] >= 1
    assert all(image_count >= 2 for image_count in rendered_page_image_counts[1:])
    assert page_two_footer_blocks
    assert metrics["text_retention_ratio"] == 1.0
    assert FIXTURE.read_bytes() == source_before
