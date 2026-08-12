from __future__ import annotations

import json
import math
import os
import shutil
import subprocess
from pathlib import Path

import fitz
import pytest
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.conversion.docx_converter import (
    PdfToDocxConverter,
    text_retention_ratio,
)
from app.conversion.models import DocxMode


PROJECT_ROOT = Path(__file__).resolve().parents[3]
DEFAULT_REAL_PDF = (
    PROJECT_ROOT
    / "data"
    / "input"
    / "manual-docx-regression"
    / "2-ENGAGEMENT_INDIVIDUEL_ETUDIANT_2026-2027.pdf"
)
RESULT_PATH = (
    PROJECT_ROOT
    / "apps"
    / "web"
    / "test-results"
    / "docx-editable-real-document"
    / "results.json"
)
EXPECTED_TEXTS = (
    "Engagement individuel",
    "Nom de l’étudiant",
    "Comportement général",
    "Rappel de la législation française",
)
EXPECTED_SUBTITLE = "Exemplaire à remettre signé à l’administration"

pytestmark = [
    pytest.mark.docx_real_document,
    pytest.mark.regression,
]


def real_pdf_path() -> Path:
    configured = os.environ.get("QA_REAL_DOCX_PDF")
    if not configured:
        return DEFAULT_REAL_PDF
    candidate = Path(configured)
    return candidate if candidate.is_absolute() else PROJECT_ROOT / candidate


def normalized(value: str) -> str:
    return " ".join(
        value.casefold().replace("’", "'").replace("\u00a0", " ").split()
    )


def black_pixel_ratio(image: bytes) -> float:
    pixmap = fitz.Pixmap(image)
    if pixmap.colorspace is None:
        pixmap = fitz.Pixmap(fitz.csRGB, pixmap)
    samples = memoryview(pixmap.samples)
    channel_count = pixmap.n
    color_channels = channel_count - int(pixmap.alpha)
    pixel_count = pixmap.width * pixmap.height
    sample_step = max(1, pixel_count // 50_000)
    visible = 0
    black = 0
    for pixel_index in range(0, pixel_count, sample_step):
        offset = pixel_index * channel_count
        if pixmap.alpha and samples[offset + color_channels] <= 20:
            continue
        visible += 1
        if all(samples[offset + channel] <= 32 for channel in range(color_channels)):
            black += 1
    return black / visible if visible else 0.0


def section_word_counts(document: Document) -> list[int]:
    counts: list[int] = []
    current_count = 0
    for element in document.element.body.iterchildren():
        text = " ".join(
            node.text or ""
            for node in element.xpath(".//w:t")
        )
        current_count += len(text.split())
        if (
            element.tag == qn("w:p")
            and element.xpath("./w:pPr/w:sectPr")
        ):
            counts.append(current_count)
            current_count = 0
    counts.append(current_count)
    return counts


def section_paragraph_counts(document: Document) -> list[int]:
    counts: list[int] = []
    current_count = 0
    for element in document.element.body.iterchildren():
        if element.tag == qn("w:p"):
            text = " ".join(
                node.text or ""
                for node in element.xpath(".//w:t")
            )
            if text.strip():
                current_count += 1
            if element.xpath("./w:pPr/w:sectPr"):
                counts.append(current_count)
                current_count = 0
    counts.append(current_count)
    return counts


def spacing_value(paragraph: object) -> float | None:
    value = paragraph.paragraph_format.line_spacing
    return float(value) if isinstance(value, float) else None


def point_value(value: object) -> float:
    return float(value.pt) if value is not None else 0


def estimated_vertical_fill_ratios(
    document: Document,
    word_counts: list[int],
    paragraph_counts: list[int],
) -> list[float]:
    ratios: list[float] = []
    for index, word_count in enumerate(word_counts):
        section = document.sections[min(index, len(document.sections) - 1)]
        usable_height = (
            section.page_height.pt
            - section.top_margin.pt
            - section.bottom_margin.pt
        )
        estimated_lines = max(1, math.ceil(word_count / 12))
        estimated_height = (
            estimated_lines * 9 * 1.12
            + paragraph_counts[index] * 2.5
        )
        ratios.append(round(estimated_height / usable_height, 4))
    return ratios


def libreoffice_page_metrics(
    docx_path: Path,
    temporary_directory: Path,
) -> tuple[int | None, list[int], str | None]:
    executable = shutil.which("libreoffice")
    if executable is None:
        return None, [], "LibreOffice indisponible"
    profile = temporary_directory / "libreoffice-profile"
    cache = temporary_directory / "xdg-cache"
    config = temporary_directory / "xdg-config"
    runtime = temporary_directory / "xdg-runtime"
    for directory in (cache, config, runtime):
        directory.mkdir(mode=0o700)
    environment = {
        **os.environ,
        "XDG_CACHE_HOME": str(cache),
        "XDG_CONFIG_HOME": str(config),
        "XDG_RUNTIME_DIR": str(runtime),
    }
    try:
        completed = subprocess.run(
            [
                executable,
                "--headless",
                f"-env:UserInstallation={profile.as_uri()}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(temporary_directory),
                str(docx_path),
            ],
            check=False,
            capture_output=True,
            text=True,
            timeout=60,
            env=environment,
        )
    except (OSError, subprocess.TimeoutExpired):
        return None, [], "Rendu LibreOffice indisponible dans cet environnement"
    rendered_pdf = temporary_directory / f"{docx_path.stem}.pdf"
    if completed.returncode != 0 or not rendered_pdf.is_file():
        return None, [], "Rendu LibreOffice bloqué dans cet environnement"
    with fitz.open(rendered_pdf) as rendered:
        page_word_counts = [
            len(page.get_text("words"))
            for page in rendered
        ]
        return len(rendered), page_word_counts, None


def test_real_pdf_produces_an_editable_docx(tmp_path: Path) -> None:
    source_path = real_pdf_path()
    if not source_path.is_file():
        pytest.skip(
            "PDF réel absent. Définissez QA_REAL_DOCX_PDF pour activer ce test."
        )

    output_path = tmp_path / "editable.docx"
    converter = PdfToDocxConverter()
    artifact = converter.convert(
        source_path,
        output_path,
        mode=DocxMode.EDITABLE,
    )
    document = Document(output_path)
    with fitz.open(source_path) as source:
        source_page_count = len(source)
        source_text = "\n".join(
            page.get_text("text", sort=True) for page in source
        )
    docx_text = "\n".join(
        [
            *(paragraph.text for paragraph in document.paragraphs),
            *(
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ),
        ]
    )
    retention_ratio = text_retention_ratio(source_text, docx_text)
    meaningful_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    list_count = sum(
        paragraph.style.name in {"List Bullet", "List Number"}
        for paragraph in meaningful_paragraphs
    )
    highlight_present = any(
        run.font.highlight_color == WD_COLOR_INDEX.YELLOW
        for paragraph in meaningful_paragraphs
        if "nom de l" in normalized(paragraph.text)
        for run in paragraph.runs
    )
    border_present = any(
        paragraph._p.xpath(".//w:pBdr")
        for paragraph in meaningful_paragraphs
    ) or bool(document.tables)
    image_parts = [
        part
        for part in document.part.package.parts
        if part.content_type.startswith("image/")
    ]
    black_pixel_ratios = [
        black_pixel_ratio(part.blob)
        for part in image_parts
    ]
    logo_background_acceptable = (
        bool(black_pixel_ratios) and min(black_pixel_ratios) < 0.15
    )
    normalized_docx_text = normalized(docx_text)
    title_presence = {
        expected: normalized(expected) in normalized_docx_text
        for expected in EXPECTED_TEXTS
    }
    centered_paragraphs = [
        paragraph
        for paragraph in meaningful_paragraphs
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    ]
    long_centered_paragraphs = [
        paragraph
        for paragraph in centered_paragraphs
        if len(paragraph.text.split()) >= 12
    ]
    left_paragraphs = [
        paragraph
        for paragraph in meaningful_paragraphs
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT
    ]
    right_paragraphs = [
        paragraph
        for paragraph in meaningful_paragraphs
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    ]
    justified_paragraphs = [
        paragraph
        for paragraph in meaningful_paragraphs
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    ]
    geometry_indent_paragraphs = [
        paragraph
        for paragraph in meaningful_paragraphs
        if any(
            abs(point_value(indent)) >= 2
            for indent in (
                paragraph.paragraph_format.left_indent,
                paragraph.paragraph_format.right_indent,
                paragraph.paragraph_format.first_line_indent,
            )
        )
    ]
    fully_bold_paragraphs = []
    mixed_bold_paragraphs = []
    for paragraph in meaningful_paragraphs:
        text_runs = [run for run in paragraph.runs if run.text.strip()]
        if text_runs and all(run.bold is True for run in text_runs):
            fully_bold_paragraphs.append(paragraph)
        if (
            any(run.bold is True for run in text_runs)
            and any(run.bold is False for run in text_runs)
        ):
            mixed_bold_paragraphs.append(paragraph)
    empty_bullets = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.style.name in {"List Bullet", "List Number"}
        and not paragraph.text.strip()
    ]
    title_paragraphs = [
        paragraph
        for paragraph in meaningful_paragraphs
        if normalized(paragraph.text) == normalized("Engagement individuel")
    ]
    subtitle_paragraphs = [
        paragraph
        for paragraph in meaningful_paragraphs
        if normalized(paragraph.text) == normalized(EXPECTED_SUBTITLE)
    ]
    drawing_paragraph_indexes = [
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph._p.xpath(".//w:drawing")
    ]
    header_drawing_count = sum(
        len(section.header._element.xpath(".//w:drawing"))
        for section in document.sections
    )
    title_paragraph_index = next(
        (
            index
            for index, paragraph in enumerate(document.paragraphs)
            if normalized(paragraph.text) == normalized("Engagement individuel")
        ),
        None,
    )
    structural_page_word_counts = section_word_counts(document)
    structural_page_paragraph_counts = section_paragraph_counts(document)
    vertical_fill_ratios = estimated_vertical_fill_ratios(
        document,
        structural_page_word_counts,
        structural_page_paragraph_counts,
    )
    structural_quasi_empty_pages = sum(
        word_count < 15
        for word_count in structural_page_word_counts
    )
    rendered_page_count, rendered_page_word_counts, render_reason = (
        libreoffice_page_metrics(output_path, tmp_path)
    )
    rendered_quasi_empty_pages = sum(
        word_count < 15
        for word_count in rendered_page_word_counts
    )
    explicit_page_breaks = len(
        document.element.xpath(".//w:br[@w:type='page']")
    )
    exact_line_rule_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if any(
            spacing.get(qn("w:lineRule")) == "exact"
            for spacing in paragraph._p.xpath("./w:pPr/w:spacing")
        )
    ]
    line_spacing_values = [
        value
        for paragraph in meaningful_paragraphs
        if (value := spacing_value(paragraph)) is not None
    ]
    long_paragraphs = [
        paragraph
        for paragraph in meaningful_paragraphs
        if paragraph.style.name == "Normal"
        if len(paragraph.text.split()) >= 30
    ]
    long_line_spacing_values = [
        value
        for paragraph in long_paragraphs
        if (value := spacing_value(paragraph)) is not None
    ]
    ordinary_paragraphs = [
        paragraph
        for paragraph in meaningful_paragraphs
        if paragraph.style.name == "Normal"
        and len(paragraph.text.split()) < 30
        and not paragraph._p.xpath(".//w:pBdr")
    ]
    ordinary_line_spacing_values = [
        value
        for paragraph in ordinary_paragraphs
        if (value := spacing_value(paragraph)) is not None
    ]
    list_paragraphs = [
        paragraph
        for paragraph in meaningful_paragraphs
        if paragraph.style.name in {"List Bullet", "List Number"}
    ]
    list_line_spacing_values = [
        value
        for paragraph in list_paragraphs
        if (value := spacing_value(paragraph)) is not None
    ]
    paragraph_space_before_values = [
        point_value(paragraph.paragraph_format.space_before)
        for paragraph in meaningful_paragraphs
    ]
    paragraph_space_after_values = [
        point_value(paragraph.paragraph_format.space_after)
        for paragraph in meaningful_paragraphs
    ]
    border_paragraphs = [
        paragraph
        for paragraph in meaningful_paragraphs
        if paragraph._p.xpath(".//w:pBdr")
    ]
    border_padding_values = [
        int(edge.get(qn("w:space"), "0"))
        for paragraph in border_paragraphs
        for edge in paragraph._p.xpath("./w:pPr/w:pBdr/*")
    ]
    dense_section_count = sum(
        ratio > 1.05
        for ratio in vertical_fill_ratios
    )
    estimated_line_counts = [
        max(1, math.ceil(len(paragraph.text.split()) / 12))
        for paragraph in meaningful_paragraphs
    ]
    average_words_per_paragraph = (
        len(docx_text.split()) / len(meaningful_paragraphs)
    )
    average_estimated_lines_per_paragraph = (
        sum(estimated_line_counts) / len(estimated_line_counts)
    )
    checks = {
        "openable": output_path.stat().st_size > 0,
        "editableText": len(docx_text.split()) >= 100,
        "retention": retention_ratio >= 0.95,
        "paragraphCount": len(meaningful_paragraphs) >= 30,
        "notImageOnly": len(meaningful_paragraphs) > len(image_parts),
        "reasonableImages": (
            1 <= len(image_parts) <= source_page_count * 4
        ),
        "logoBackground": logo_background_acceptable,
        "reasonablePagination": len(document.sections) == source_page_count,
        "titles": all(title_presence.values()),
        "titleAndSubtitleSeparated": (
            bool(title_paragraphs) and bool(subtitle_paragraphs)
        ),
        "logoBeforeTitle": (
            header_drawing_count > 0
            or (
                bool(drawing_paragraph_indexes)
                and title_paragraph_index is not None
                and drawing_paragraph_indexes[0] < title_paragraph_index
            )
        ),
        "lists": list_count > 0,
        "emptyBullets": len(empty_bullets) == 0,
        "longCenteredParagraphs": len(long_centered_paragraphs) == 0,
        "paragraphGeometry": (
            bool(left_paragraphs)
            and bool(justified_paragraphs)
            and bool(geometry_indent_paragraphs)
        ),
        "flowCanonicalization": (
            converter.last_flow_metrics["same_baseline_fragments_merged"] > 0
        ),
        "mixedBoldRuns": len(mixed_bold_paragraphs) > 0,
        "structuralQuasiEmptyPages": structural_quasi_empty_pages == 0,
        "pageInflation": len(document.sections) <= source_page_count + 1,
        "explicitPageBreaks": explicit_page_breaks <= source_page_count - 1,
        "safeLineRules": len(exact_line_rule_paragraphs) == 0,
        "readableLongLineSpacing": (
            bool(long_line_spacing_values)
            and min(long_line_spacing_values) >= 1.12
            and max(long_line_spacing_values) <= 1.2
        ),
        "readableOrdinaryLineSpacing": (
            bool(ordinary_line_spacing_values)
            and min(ordinary_line_spacing_values) >= 1.08
            and max(ordinary_line_spacing_values) <= 1.2
        ),
        "readableListLineSpacing": (
            bool(list_line_spacing_values)
            and min(list_line_spacing_values) >= 1.05
            and max(list_line_spacing_values) <= 1.12
        ),
        "reasonableParagraphSpacing": (
            1.5
            <= (
                sum(paragraph_space_after_values)
                / len(paragraph_space_after_values)
            )
            <= 4
            and max(paragraph_space_before_values, default=0) <= 8
            and max(paragraph_space_after_values, default=0) <= 6
        ),
        "borderPadding": (
            bool(border_padding_values)
            and min(border_padding_values) >= 7
            and all(
                point_value(paragraph.paragraph_format.space_before) >= 3
                and point_value(paragraph.paragraph_format.space_after) >= 3
                for paragraph in border_paragraphs
            )
        ),
        "sectionDensity": dense_section_count == 0,
        "renderedPageInflation": (
            rendered_page_count is None
            or rendered_page_count <= source_page_count + 1
        ),
        "renderedQuasiEmptyPages": (
            rendered_page_count is None
            or rendered_quasi_empty_pages == 0
        ),
        "highlight": highlight_present,
        "border": border_present,
    }
    RESULT_PATH.parent.mkdir(parents=True, exist_ok=True)
    RESULT_PATH.write_text(
        json.dumps(
            {
                "file": source_path.name,
                "docxMode": "editable",
                "sourcePageCount": source_page_count,
                "pageCountDelta": (
                    rendered_page_count - source_page_count
                    if rendered_page_count is not None
                    else len(document.sections) - source_page_count
                ),
                "sameBaselineFragmentsMerged": converter.last_flow_metrics[
                    "same_baseline_fragments_merged"
                ],
                "rawTextFragments": converter.last_flow_metrics[
                    "raw_text_fragments"
                ],
                "logicalLineCount": converter.last_flow_metrics[
                    "logical_lines"
                ],
                "sourceTextCharacters": len(source_text),
                "docxTextCharacters": len(docx_text),
                "textRetentionRatio": round(retention_ratio, 4),
                "paragraphCount": len(meaningful_paragraphs),
                "editableWordCount": len(docx_text.split()),
                "averageWordsPerParagraph": round(
                    average_words_per_paragraph,
                    2,
                ),
                "averageEstimatedLinesPerParagraph": round(
                    average_estimated_lines_per_paragraph,
                    2,
                ),
                "imageCount": len(image_parts),
                "bodyInlineImageCount": len(document.inline_shapes),
                "headerDrawingCount": header_drawing_count,
                "logoPresent": len(image_parts) > 0,
                "logoBackgroundAcceptable": logo_background_acceptable,
                "imageBlackPixelRatios": [
                    round(ratio, 4) for ratio in black_pixel_ratios
                ],
                "docxSectionCount": len(document.sections),
                "titlePresence": title_presence,
                "listCount": list_count,
                "centeredParagraphCount": len(centered_paragraphs),
                "longCenteredParagraphCount": len(long_centered_paragraphs),
                "leftParagraphCount": len(left_paragraphs),
                "rightParagraphCount": len(right_paragraphs),
                "justifiedParagraphCount": len(justified_paragraphs),
                "geometryIndentParagraphCount": len(
                    geometry_indent_paragraphs
                ),
                "fullyBoldParagraphCount": len(fully_bold_paragraphs),
                "mixedBoldParagraphCount": len(mixed_bold_paragraphs),
                "emptyBulletCount": len(empty_bullets),
                "structuralPageWordCounts": structural_page_word_counts,
                "structuralPageParagraphCounts": (
                    structural_page_paragraph_counts
                ),
                "estimatedVerticalFillRatios": vertical_fill_ratios,
                "denseSectionCount": dense_section_count,
                "estimatedPageCount": len(document.sections),
                "renderedPageCount": rendered_page_count,
                "renderedPageWordCounts": rendered_page_word_counts,
                "quasiEmptyPageCount": (
                    rendered_quasi_empty_pages
                    if rendered_page_count is not None
                    else structural_quasi_empty_pages
                ),
                "explicitPageBreakCount": explicit_page_breaks,
                "exactLineRuleParagraphCount": len(
                    exact_line_rule_paragraphs
                ),
                "lineSpacing": {
                    "minimum": round(min(line_spacing_values), 4),
                    "maximum": round(max(line_spacing_values), 4),
                    "average": round(
                        sum(line_spacing_values) / len(line_spacing_values),
                        4,
                    ),
                    "longParagraphMinimum": round(
                        min(long_line_spacing_values),
                        4,
                    ),
                    "ordinaryParagraphMinimum": round(
                        min(ordinary_line_spacing_values),
                        4,
                    ),
                    "listMinimum": round(
                        min(list_line_spacing_values),
                        4,
                    ),
                },
                "paragraphSpacingPoints": {
                    "beforeMaximum": round(
                        max(paragraph_space_before_values),
                        2,
                    ),
                    "afterMinimum": round(
                        min(paragraph_space_after_values),
                        2,
                    ),
                    "afterMaximum": round(
                        max(paragraph_space_after_values),
                        2,
                    ),
                    "afterAverage": round(
                        sum(paragraph_space_after_values)
                        / len(paragraph_space_after_values),
                        2,
                    ),
                },
                "borderPaddingPoints": border_padding_values,
                "pageMeasurement": (
                    "libreoffice"
                    if rendered_page_count is not None
                    else "sections-docx"
                ),
                "renderWarning": render_reason,
                "readabilityStatus": (
                    "KO"
                    if not all(checks.values())
                    else "R"
                    if rendered_page_count is None
                    else "OK"
                ),
                "highlightPresent": highlight_present,
                "borderPresent": border_present,
                "warnings": list(artifact.warnings),
                "status": "passed" if all(checks.values()) else "failed",
                "checks": checks,
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )

    assert all(checks.values()), checks
