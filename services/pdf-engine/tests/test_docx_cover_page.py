from __future__ import annotations

import json
import math
import os
import re
import shutil
import subprocess
from pathlib import Path
from typing import Any

import fitz
import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.conversion.docx_converter import (
    PdfToDocxConverter,
    text_retention_ratio,
)
from app.conversion.models import DocxMode


PROJECT_ROOT = Path(__file__).resolve().parents[3]
DEFAULT_COVER_PDF = (
    PROJECT_ROOT
    / "data"
    / "input"
    / "manual-docx-regression"
    / "cort_test.pdf"
)
CONTROL_OUTPUT = (
    PROJECT_ROOT
    / "data"
    / "output"
    / "cort_test-cover-page-editable.docx"
)
COVER_RESULTS_PATH = (
    PROJECT_ROOT
    / "apps"
    / "web"
    / "test-results"
    / "docx-cover-page"
    / "results.json"
)
BLUE = (0.05, 0.25, 0.55)


def cover_pdf_path() -> Path:
    configured = os.environ.get("QA_COVER_PAGE_PDF")
    if not configured:
        return DEFAULT_COVER_PDF
    candidate = Path(configured)
    return candidate if candidate.is_absolute() else PROJECT_ROOT / candidate


def make_logo_png() -> bytes:
    logo_pdf = fitz.open()
    page = logo_pdf.new_page(width=180, height=60)
    page.draw_rect(
        fitz.Rect(0, 0, 180, 60),
        color=BLUE,
        fill=(0.92, 0.96, 1),
        width=1,
    )
    page.insert_text(
        (18, 39),
        "EXEMPLE",
        fontsize=24,
        fontname="hebo",
        color=BLUE,
    )
    image = page.get_pixmap(dpi=144, alpha=False).tobytes("png")
    logo_pdf.close()
    return image


def make_synthetic_cover_pdf() -> bytes:
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    page.insert_image(
        fitz.Rect(45, 42, 150, 77),
        stream=make_logo_png(),
    )
    page.insert_textbox(
        fitz.Rect(80, 275, 515, 345),
        "RAPPORT",
        fontsize=30,
        fontname="hebo",
        color=BLUE,
        align=fitz.TEXT_ALIGN_CENTER,
    )
    page.insert_textbox(
        fitz.Rect(85, 350, 510, 410),
        "Analyse synthétique des cultures\nCampagne de référence",
        fontsize=17,
        fontname="heit",
        color=BLUE,
        align=fitz.TEXT_ALIGN_CENTER,
    )
    page.insert_textbox(
        fitz.Rect(120, 545, 475, 575),
        "Juillet-Août 2023",
        fontsize=14,
        fontname="helv",
        color=BLUE,
        align=fitz.TEXT_ALIGN_CENTER,
    )
    page.insert_textbox(
        fitz.Rect(80, 720, 515, 750),
        "MARTIN Camille - DUPONT Alex",
        fontsize=12,
        fontname="helv",
        color=BLUE,
        align=fitz.TEXT_ALIGN_CENTER,
    )
    content = document.tobytes(garbage=4, deflate=True)
    document.close()
    return content


def paragraph_position_ratios(document: Document) -> dict[str, float]:
    section = document.sections[0]
    page_height = section.page_height.pt
    cursor = section.top_margin.pt
    positions: dict[str, float] = {}
    for paragraph in document.paragraphs:
        before = paragraph.paragraph_format.space_before
        cursor += before.pt if before is not None else 0
        text = " ".join(paragraph.text.split())
        if text:
            positions[text] = cursor / page_height
        drawing_heights = [
            int(extent.get("cy", "0")) / 12700
            for extent in paragraph._p.xpath(".//wp:extent")
        ]
        if drawing_heights:
            content_height = max(drawing_heights)
        else:
            font_sizes = [
                run.font.size.pt
                for run in paragraph.runs
                if run.font.size is not None and run.text.strip()
            ]
            font_size = max(font_sizes, default=11)
            line_count = max(1, math.ceil(len(text) / 55))
            spacing = paragraph.paragraph_format.line_spacing
            line_spacing = float(spacing) if isinstance(spacing, float) else 1.1
            content_height = font_size * line_spacing * line_count
        cursor += content_height
        after = paragraph.paragraph_format.space_after
        cursor += after.pt if after is not None else 0
    return positions


def rendered_page_count(docx_path: Path, output_directory: Path) -> int | None:
    executable = shutil.which("libreoffice")
    if executable is None:
        return None
    profile = output_directory / "lo-profile"
    runtime = output_directory / "lo-runtime"
    runtime.mkdir(mode=0o700)
    try:
        result = subprocess.run(
            [
                executable,
                "--headless",
                f"-env:UserInstallation={profile.as_uri()}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_directory),
                str(docx_path),
            ],
            check=False,
            capture_output=True,
            text=True,
            timeout=45,
            env={
                **os.environ,
                "XDG_RUNTIME_DIR": str(runtime),
            },
        )
    except (OSError, subprocess.TimeoutExpired):
        return None
    rendered_pdf = output_directory / f"{docx_path.stem}.pdf"
    if result.returncode != 0 or not rendered_pdf.is_file():
        return None
    with fitz.open(rendered_pdf) as rendered:
        return len(rendered)


def inspect_cover_conversion(
    source_path: Path,
    output_path: Path,
    temporary_directory: Path,
) -> dict[str, Any]:
    converter = PdfToDocxConverter()
    converter.convert(source_path, output_path, mode=DocxMode.EDITABLE)
    document = Document(output_path)
    with fitz.open(source_path) as source:
        page = source[0]
        blocks = page.get_text("dict", sort=True).get("blocks", [])
        analysis = converter._analyze_cover_page(
            page,
            blocks,
            allow_cover_page=True,
        )
        source_text = page.get_text("text", sort=True)
    paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    docx_text = "\n".join(paragraph.text for paragraph in paragraphs)
    positions = paragraph_position_ratios(document)
    title = next(
        paragraph
        for paragraph in paragraphs
        if paragraph.text.strip() == "RAPPORT"
    )
    author = next(
        paragraph
        for paragraph in paragraphs
        if re.search(r"MARTIN|ROLLAND", paragraph.text, re.IGNORECASE)
    )
    date = next(
        paragraph
        for paragraph in paragraphs
        if re.search(r"20\d{2}", paragraph.text)
    )
    title_sizes = [
        run.font.size.pt
        for run in title.runs
        if run.font.size is not None
    ]
    secondary_sizes = [
        run.font.size.pt
        for paragraph in paragraphs
        if paragraph is not title
        for run in paragraph.runs
        if run.font.size is not None
    ]
    explicit_page_breaks = document.element.xpath(".//w:br[@w:type='page']")
    image_only = not docx_text.strip() or len(document.inline_shapes) >= len(paragraphs)
    retention = text_retention_ratio(source_text, docx_text)
    rendered_pages = rendered_page_count(output_path, temporary_directory)
    title_ratio = positions[title.text.strip()]
    date_ratio = positions[date.text.strip()]
    author_ratio = positions[author.text.strip()]
    checks = {
        "detectedAsCover": analysis.is_cover_page,
        "editableText": len(docx_text.split()) >= 10,
        "notImageOnly": not image_only,
        "singleSection": len(document.sections) == 1,
        "singleRenderedPage": rendered_pages in {None, 1},
        "noExplicitPageBreak": not explicit_page_breaks,
        "titleCentered": title.alignment == WD_ALIGN_PARAGRAPH.CENTER,
        "titleDominant": (
            bool(title_sizes)
            and bool(secondary_sizes)
            and max(title_sizes) >= max(secondary_sizes) * 1.35
        ),
        "titleCentral": 0.25 <= title_ratio <= 0.6,
        "dateBelowTitle": date_ratio >= title_ratio + 0.15,
        "authorNearBottom": author_ratio >= 0.7,
        "textRetention": retention >= 0.95,
        "logoBeforeTitle": next(
            index
            for index, paragraph in enumerate(document.paragraphs)
            if paragraph._p.xpath(".//w:drawing")
        )
        < next(
            index
            for index, paragraph in enumerate(document.paragraphs)
            if paragraph.text.strip() == "RAPPORT"
        ),
    }
    return {
        "status": "passed" if all(checks.values()) else "failed",
        "file": source_path.name,
        "layoutType": "cover_page" if analysis.is_cover_page else "normal",
        "sourceTextBlockCount": analysis.text_block_count,
        "docxParagraphCount": len(paragraphs),
        "sourceWhitespaceRatio": analysis.whitespace_ratio,
        "sourceTitlePositionRatio": analysis.title_position_ratio,
        "docxTitlePositionRatio": round(title_ratio, 4),
        "docxDatePositionRatio": round(date_ratio, 4),
        "docxAuthorPositionRatio": round(author_ratio, 4),
        "titleSizeRatio": round(max(title_sizes) / max(secondary_sizes), 4),
        "textRetentionRatio": round(retention, 4),
        "imageOnly": image_only,
        "renderedPageCount": rendered_pages,
        "renderStatus": (
            "available" if rendered_pages is not None else "unavailable"
        ),
        "checks": checks,
    }


def write_cover_results(metrics: dict[str, Any]) -> None:
    COVER_RESULTS_PATH.parent.mkdir(parents=True, exist_ok=True)
    COVER_RESULTS_PATH.write_text(
        json.dumps(metrics, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def test_cover_strategy_is_limited_to_first_sparse_presentation_page(
    tmp_path: Path,
) -> None:
    cover_path = tmp_path / "cover.pdf"
    cover_path.write_bytes(make_synthetic_cover_pdf())
    converter = PdfToDocxConverter()
    with fitz.open(cover_path) as source:
        page = source[0]
        blocks = page.get_text("dict", sort=True).get("blocks", [])
        disabled = converter._analyze_cover_page(
            page,
            blocks,
            allow_cover_page=False,
        )
    assert disabled.is_cover_page is False

    content = fitz.open()
    content_page = content.new_page(width=595, height=842)
    for index in range(10):
        content_page.insert_textbox(
            fitz.Rect(55, 55 + index * 65, 540, 105 + index * 65),
            (
                f"Paragraphe {index + 1} de contenu ordinaire avec assez de "
                "texte pour représenter une page structurée."
            ),
            fontsize=11,
            fontname="helv",
        )
    content_path = tmp_path / "content.pdf"
    content.save(content_path)
    content.close()
    with fitz.open(content_path) as source:
        page = source[0]
        blocks = page.get_text("dict", sort=True).get("blocks", [])
        analysis = converter._analyze_cover_page(
            page,
            blocks,
            allow_cover_page=True,
        )
    assert analysis.is_cover_page is False


@pytest.mark.regression
def test_synthetic_cover_page_preserves_editable_vertical_layout(
    tmp_path: Path,
) -> None:
    source_path = tmp_path / "conversion-cover-page.pdf"
    output_path = tmp_path / "conversion-cover-page.docx"
    source_path.write_bytes(make_synthetic_cover_pdf())

    metrics = inspect_cover_conversion(source_path, output_path, tmp_path)
    write_cover_results(metrics)

    assert metrics["status"] == "passed", metrics["checks"]


@pytest.mark.docx_cover_page
@pytest.mark.regression
def test_local_cover_page_regression(tmp_path: Path) -> None:
    source_path = cover_pdf_path()
    if not source_path.is_file():
        pytest.skip(
            "PDF de couverture absent. Définissez QA_COVER_PAGE_PDF pour "
            "activer ce test."
        )

    CONTROL_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    metrics = inspect_cover_conversion(source_path, CONTROL_OUTPUT, tmp_path)
    write_cover_results(metrics)

    assert metrics["status"] == "passed", metrics["checks"]
