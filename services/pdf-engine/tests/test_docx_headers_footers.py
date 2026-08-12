from __future__ import annotations

import asyncio
import os
import shutil
import subprocess
from pathlib import Path
from typing import Any, Literal

import fitz
import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.conversion.docx_converter import PdfToDocxConverter
from app.conversion.header_footer import (
    PageLayoutInput,
    detect_document_layout,
    parse_page_number,
)
from app.conversion.models import ConversionOptions, OcrMode, TargetFormat
from app.conversion.service import prepare_searchable_pdf


PAGE_WIDTH = 612
PAGE_HEIGHT = 792
OCR_BINARIES = ("ocrmypdf", "tesseract", "gs", "qpdf")
OCR_BINARIES_AVAILABLE = all(shutil.which(binary) for binary in OCR_BINARIES)
LIBREOFFICE_AVAILABLE = shutil.which("libreoffice") is not None


def _insert_edge_text(
    page: fitz.Page,
    text: str,
    *,
    y: float,
    alignment: Literal["left", "center", "right"] = "center",
) -> None:
    word_alignment = {
        "left": fitz.TEXT_ALIGN_LEFT,
        "center": fitz.TEXT_ALIGN_CENTER,
        "right": fitz.TEXT_ALIGN_RIGHT,
    }[alignment]
    page.insert_textbox(
        fitz.Rect(48, y, PAGE_WIDTH - 48, y + 24),
        text,
        fontsize=10,
        fontname="helv",
        align=word_alignment,
    )


def _make_layout_pdf(
    pages: list[dict[str, str | None]],
    *,
    footer_alignment: Literal["left", "center", "right"] = "center",
) -> bytes:
    document = fitz.open()
    for page_index, page_content in enumerate(pages):
        page = document.new_page(width=PAGE_WIDTH, height=PAGE_HEIGHT)
        header = page_content.get("header")
        if header:
            _insert_edge_text(page, header, y=28)
        page.insert_textbox(
            fitz.Rect(60, 170, PAGE_WIDTH - 60, 600),
            page_content.get("body") or f"Contenu principal page {page_index + 1}",
            fontsize=12,
            fontname="helv",
        )
        footer = page_content.get("footer")
        if footer:
            _insert_edge_text(
                page,
                footer,
                y=746,
                alignment=footer_alignment,
            )
    content = document.tobytes(garbage=4, deflate=True)
    document.close()
    return content


def _convert(tmp_path: Path, content: bytes) -> Document:
    source = tmp_path / "source.pdf"
    output = tmp_path / "output.docx"
    source.write_bytes(content)
    PdfToDocxConverter().convert(source, output)
    return Document(output)


def _field_instructions(story: Any) -> list[str]:
    return [
        (element.text or "").strip()
        for element in story._element.xpath(".//w:instrText")
    ]


def _story_text(story: Any) -> str:
    return "\n".join(paragraph.text for paragraph in story.paragraphs).strip()


def _body_text(document: Document) -> str:
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _source_block(
    text: str,
    bbox: tuple[float, float, float, float],
) -> dict[str, Any]:
    return {
        "type": 0,
        "bbox": bbox,
        "lines": [
            {
                "bbox": bbox,
                "spans": [
                    {
                        "text": text,
                        "bbox": bbox,
                        "font": "Arial",
                        "size": 10,
                        "flags": 0,
                        "color": 0,
                    }
                ],
            }
        ],
    }


def test_repeated_header_and_numeric_pagination_become_word_stories(
    tmp_path: Path,
) -> None:
    document = _convert(
        tmp_path,
        _make_layout_pdf(
            [
                {
                    "header": "HEADER IDENTIQUE",
                    "body": f"Corps différent {page_number}",
                    "footer": str(page_number),
                }
                for page_number in range(1, 4)
            ]
        ),
    )

    assert len(document.sections) == 3
    assert all(
        _story_text(section.header) == "HEADER IDENTIQUE"
        for section in document.sections
    )
    assert all(
        _field_instructions(section.footer) == ["PAGE"]
        for section in document.sections
    )
    assert all(
        section.footer.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
        for section in document.sections
    )
    body = _body_text(document)
    assert "HEADER IDENTIQUE" not in body
    assert all(f"Corps différent {page_number}" in body for page_number in range(1, 4))
    assert not any(line.strip() in {"1", "2", "3"} for line in body.splitlines())


def test_page_x_sur_y_uses_page_and_numpages_fields(tmp_path: Path) -> None:
    document = _convert(
        tmp_path,
        _make_layout_pdf(
            [
                {
                    "header": "RAPPORT ANNUEL",
                    "body": f"Section {page_number}",
                    "footer": f"Page {page_number} sur 3",
                }
                for page_number in range(1, 4)
            ]
        ),
    )

    for section in document.sections:
        assert _field_instructions(section.footer) == ["PAGE", "NUMPAGES"]
        assert "Page " in _story_text(section.footer)
        assert " sur " in _story_text(section.footer)
    body = _body_text(document)
    assert "Page 1 sur 3" not in body
    assert "Page 2 sur 3" not in body
    assert "Page 3 sur 3" not in body


def test_repeated_right_aligned_footer_keeps_alignment(tmp_path: Path) -> None:
    document = _convert(
        tmp_path,
        _make_layout_pdf(
            [
                {
                    "header": None,
                    "body": f"Contenu {page_number}",
                    "footer": "CONFIDENTIEL",
                }
                for page_number in range(1, 4)
            ],
            footer_alignment="right",
        ),
    )

    assert all(
        _story_text(section.footer) == "CONFIDENTIEL"
        for section in document.sections
    )
    assert all(
        section.footer.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
        for section in document.sections
    )
    assert all(
        section.footer.paragraphs[0].paragraph_format.line_spacing
        == pytest.approx(1.1)
        for section in document.sections
    )


def test_same_baseline_header_fragments_are_one_logical_story_line() -> None:
    inputs = tuple(
        PageLayoutInput(
            page_index=page_index,
            width=PAGE_WIDTH,
            height=PAGE_HEIGHT,
            blocks=(
                _source_block("HEADER", (48, 28, 102, 42)),
                _source_block("IDENTIQUE", (112, 28, 180, 42)),
                _source_block(
                    f"Corps {page_index + 1}",
                    (60, 180, 300, 200),
                ),
            ),
        )
        for page_index in range(3)
    )
    layout = detect_document_layout(inputs)
    output = Document()
    converter = PdfToDocxConverter()
    converter.last_flow_metrics = {
        "raw_text_fragments": 0,
        "logical_lines": 0,
        "same_baseline_fragments_merged": 0,
        "paragraphs": 0,
    }
    source = fitz.open()
    page = source.new_page(width=PAGE_WIDTH, height=PAGE_HEIGHT)
    try:
        converter._configure_header_footer(
            output.sections[0],
            page,
            layout.pages[0],
        )
    finally:
        source.close()

    header_paragraphs = [
        paragraph
        for paragraph in output.sections[0].header.paragraphs
        if paragraph.text.strip()
    ]
    assert len(header_paragraphs) == 1
    assert header_paragraphs[0].text == "HEADER IDENTIQUE"
    assert converter.last_flow_metrics["same_baseline_fragments_merged"] == 1
    assert all(
        block.story_type == "header"
        and block.classification_confidence == 1
        for block in layout.pages[0].header_candidates
    )
    assert all(block.story_type == "body" for block in layout.pages[0].body_blocks)


def _make_repeated_logo_pdf() -> bytes:
    logo = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 120, 24), False)
    logo.clear_with(80)
    document = fitz.open()
    for page_number in range(1, 4):
        page = document.new_page(width=PAGE_WIDTH, height=PAGE_HEIGHT)
        page.insert_image(
            fitz.Rect(48, 24, 168, 48),
            stream=logo.tobytes("png"),
        )
        page.insert_text(
            (60, 180),
            f"Contenu image {page_number}",
            fontsize=12,
        )
    content = document.tobytes(garbage=4, deflate=True)
    document.close()
    return content


def test_repeated_top_image_becomes_header_without_body_duplication(
    tmp_path: Path,
) -> None:
    document = _convert(tmp_path, _make_repeated_logo_pdf())

    assert len(document.inline_shapes) == 0
    assert all(
        len(section.header._element.xpath(".//w:drawing")) == 1
        for section in document.sections
    )
    assert all(
        len(section.header.part.rels) >= 1
        for section in document.sections
    )


def test_cover_stays_empty_and_delayed_pagination_starts_at_one(
    tmp_path: Path,
) -> None:
    pages = [
        {
            "header": None,
            "body": "TITRE DE COUVERTURE\nAuteur\nParis 2026",
            "footer": None,
        },
        *[
            {
                "header": "EN-TÊTE APRÈS COUVERTURE",
                "body": f"Chapitre {page_number}",
                "footer": str(page_number),
            }
            for page_number in range(1, 4)
        ],
    ]
    document = _convert(tmp_path, _make_layout_pdf(pages))

    assert not _story_text(document.sections[0].header)
    assert not _story_text(document.sections[0].footer)
    assert all(
        _story_text(section.header) == "EN-TÊTE APRÈS COUVERTURE"
        for section in document.sections[1:]
    )
    assert all(
        _field_instructions(section.footer) == ["PAGE"]
        for section in document.sections[1:]
    )
    page_number_types = document.sections[1]._sectPr.xpath("./w:pgNumType")
    assert len(page_number_types) == 1
    assert page_number_types[0].get(qn("w:start")) == "1"


def test_ambiguous_numbers_near_bottom_stay_in_body(tmp_path: Path) -> None:
    document = _convert(
        tmp_path,
        _make_layout_pdf(
            [
                {"header": None, "body": "Texte A", "footer": "2026"},
                {"header": None, "body": "Texte B", "footer": "17"},
                {"header": None, "body": "Texte C", "footer": "42"},
            ]
        ),
    )

    assert all(not _field_instructions(section.footer) for section in document.sections)
    body = _body_text(document)
    assert all(value in body for value in ("2026", "17", "42"))


def test_header_footer_conversion_keeps_source_byte_identical(
    tmp_path: Path,
) -> None:
    content = _make_layout_pdf(
        [
            {
                "header": "SOURCE INCHANGEABLE",
                "body": f"Corps {page_number}",
                "footer": f"{page_number} / 3",
            }
            for page_number in range(1, 4)
        ]
    )
    source = tmp_path / "source.pdf"
    source.write_bytes(content)

    PdfToDocxConverter().convert(source, tmp_path / "output.docx")

    assert source.read_bytes() == content


def test_ocr_noise_uses_the_same_detector_without_losing_body_blocks() -> None:
    headers = (
        "RAPPORT CONFIDENTIEL",
        "RAPPORT CONFIDENTlEL",
        "RAPPORT  CONFIDENTIEL",
    )
    inputs = tuple(
        PageLayoutInput(
            page_index=page_index,
            width=PAGE_WIDTH,
            height=PAGE_HEIGHT,
            origin="ocr",
            blocks=(
                _source_block(header, (48, 28, 240, 42)),
                _source_block(f"Corps OCR {page_index + 1}", (60, 180, 300, 200)),
                _source_block(
                    f"Page {page_index + 1} sur 3",
                    (240, 746, 372, 760),
                ),
            ),
        )
        for page_index, header in enumerate(headers)
    )

    layout = detect_document_layout(inputs)

    assert layout.pagination is not None
    assert layout.pagination.start_number == 1
    assert all(len(page.header_candidates) == 1 for page in layout.pages)
    assert all(page.header_candidates[0].origin == "ocr" for page in layout.pages)
    assert all(len(page.body_blocks) == 1 for page in layout.pages)
    assert all(page.page_number_candidate is not None for page in layout.pages)


def test_obvious_header_changes_are_classified_in_separate_runs() -> None:
    inputs = tuple(
        PageLayoutInput(
            page_index=page_index,
            width=PAGE_WIDTH,
            height=PAGE_HEIGHT,
            blocks=(
                _source_block(
                    "Chapitre A" if page_index < 3 else "Chapitre B",
                    (48, 28, 160, 42),
                ),
                _source_block(f"Corps {page_index + 1}", (60, 180, 300, 200)),
            ),
        )
        for page_index in range(6)
    )

    layout = detect_document_layout(inputs)

    assert [page.header_candidates[0].text for page in layout.pages] == [
        "Chapitre A",
        "Chapitre A",
        "Chapitre A",
        "Chapitre B",
        "Chapitre B",
        "Chapitre B",
    ]


@pytest.mark.parametrize(
    ("value", "kind", "uses_total"),
    [
        ("3", "number", False),
        ("Page 3", "page", False),
        ("3 / 14", "slash", True),
        ("Page 3 sur 14", "page_of", True),
        ("PAGE 3 / 14", "slash", True),
    ],
)
def test_page_number_formats_are_recognized(
    value: str,
    kind: str,
    uses_total: bool,
) -> None:
    parsed = parse_page_number(value)

    assert parsed is not None
    assert parsed.kind == kind
    assert parsed.uses_total is uses_total


@pytest.mark.parametrize("value", ["2026-A", "Version 3", "15 / 3", "Page trois"])
def test_ambiguous_page_number_formats_are_rejected(value: str) -> None:
    assert parse_page_number(value) is None


def _make_scanned_layout_pdf() -> bytes:
    source = fitz.open()
    for page_number in range(1, 4):
        page = source.new_page(width=PAGE_WIDTH, height=PAGE_HEIGHT)
        page.insert_textbox(
            fitz.Rect(48, 24, PAGE_WIDTH - 48, 60),
            "REPEATED HEADER",
            fontsize=16,
            fontname="helv",
            align=fitz.TEXT_ALIGN_CENTER,
        )
        page.insert_text(
            (60, 190),
            f"BODY CONTENT PAGE {page_number}",
            fontsize=18,
            fontname="helv",
        )
        page.insert_textbox(
            fitz.Rect(48, 718, PAGE_WIDTH - 48, 760),
            f"Page {page_number}",
            fontsize=20,
            fontname="helv",
            align=fitz.TEXT_ALIGN_CENTER,
        )
    scanned = fitz.open()
    for source_page in source:
        page = scanned.new_page(width=PAGE_WIDTH, height=PAGE_HEIGHT)
        page.insert_image(
            page.rect,
            stream=source_page.get_pixmap(dpi=200, alpha=False).tobytes("png"),
        )
    content = scanned.tobytes(garbage=4, deflate=True)
    source.close()
    scanned.close()
    return content


@pytest.mark.integration
@pytest.mark.skipif(
    not OCR_BINARIES_AVAILABLE,
    reason="OCRmyPDF, Tesseract, Ghostscript et qpdf sont requis.",
)
def test_scanned_pdf_uses_production_ocr_before_header_detection(
    tmp_path: Path,
) -> None:
    source = tmp_path / "scan.pdf"
    searchable = tmp_path / "searchable.pdf"
    output = tmp_path / "output.docx"
    source.write_bytes(_make_scanned_layout_pdf())
    options = ConversionOptions(
        target_format=TargetFormat.DOCX,
        languages="eng",
        ocr_mode=OcrMode.ALWAYS,
    )

    searchable_path, ocr_used = asyncio.run(
        prepare_searchable_pdf(
            source,
            searchable,
            options=options,
            page_count=3,
            needs_ocr=True,
            temporary_directory=tmp_path,
        )
    )
    PdfToDocxConverter().convert(searchable_path, output, text_origin="ocr")
    document = Document(output)

    assert ocr_used is True
    assert all("REPEATED HEADER" in _story_text(section.header) for section in document.sections)
    assert all(_field_instructions(section.footer) == ["PAGE"] for section in document.sections)


@pytest.mark.integration
@pytest.mark.skipif(
    not LIBREOFFICE_AVAILABLE,
    reason="LibreOffice est requis pour la validation visuelle DOCX.",
)
def test_rendered_docx_places_header_and_dynamic_footer_near_page_edges(
    tmp_path: Path,
) -> None:
    source = tmp_path / "source.pdf"
    docx_output = tmp_path / "structured.docx"
    source.write_bytes(
        _make_layout_pdf(
            [
                {
                    "header": "EN-TETE VISUEL",
                    "body": f"Corps visuel {page_number}",
                    "footer": f"Page {page_number} sur 3",
                }
                for page_number in range(1, 4)
            ]
        )
    )
    PdfToDocxConverter().convert(source, docx_output)
    profile = tmp_path / "lo-profile"
    runtime = tmp_path / "lo-runtime"
    runtime.mkdir(mode=0o700)

    result = subprocess.run(
        [
            shutil.which("libreoffice") or "libreoffice",
            "--headless",
            f"-env:UserInstallation={profile.as_uri()}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(tmp_path),
            str(docx_output),
        ],
        check=False,
        capture_output=True,
        text=True,
        timeout=45,
        env={**os.environ, "XDG_RUNTIME_DIR": str(runtime)},
    )
    rendered_path = tmp_path / "structured.pdf"

    if result.returncode != 0 or not rendered_path.is_file():
        pytest.skip("Le rendu LibreOffice headless n'est pas disponible.")
    with fitz.open(rendered_path) as rendered:
        assert len(rendered) == 3
        for page_number, page in enumerate(rendered, start=1):
            blocks = [
                block
                for block in page.get_text("blocks", sort=True)
                if str(block[4]).strip()
            ]
            header = next(block for block in blocks if "EN-TETE VISUEL" in block[4])
            footer = next(
                block
                for block in blocks
                if f"Page {page_number} sur 3" in block[4]
            )
            assert header[3] <= page.rect.height * 0.15
            assert footer[1] >= page.rect.height * 0.85
