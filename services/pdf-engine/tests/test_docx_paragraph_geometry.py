from __future__ import annotations

import os
import shutil
import statistics
import subprocess
from pathlib import Path
from typing import Any

import fitz
import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.conversion.docx_converter import PdfToDocxConverter
from app.conversion.paragraph_layout import (
    ParagraphLayout,
    analyze_paragraph_layouts,
)


PAGE_RECTANGLE = fitz.Rect(0, 0, 600, 800)
CONTENT_RECTANGLE = fitz.Rect(50, 60, 550, 740)
LIBREOFFICE_AVAILABLE = shutil.which("libreoffice") is not None


def _line(
    text: str,
    *,
    left: float,
    right: float,
    baseline: float,
    font_size: float = 10,
    bold: bool = False,
) -> dict[str, Any]:
    bbox = (left, baseline - font_size, right, baseline + font_size * 0.2)
    return {
        "bbox": bbox,
        "spans": [
            {
                "text": text,
                "bbox": bbox,
                "origin": (left, baseline),
                "font": "Helvetica-Bold" if bold else "Helvetica",
                "size": font_size,
                "flags": 16 if bold else 0,
                "color": 0,
            }
        ],
    }


def _analyze(
    lines: list[dict[str, Any]],
    *,
    content_rectangle: fitz.Rect = CONTENT_RECTANGLE,
    origin: str = "native",
) -> tuple[ParagraphLayout, ...]:
    return analyze_paragraph_layouts(
        lines,
        PAGE_RECTANGLE,
        content_rectangle,
        origin=origin,  # type: ignore[arg-type]
    )


def _write_layouts(layouts: list[ParagraphLayout]) -> Document:
    document = Document()
    converter = PdfToDocxConverter()
    converter._configure_styles(document)
    source = fitz.open()
    page = source.new_page(width=600, height=800)
    try:
        for layout in layouts:
            converter._append_text_paragraph(
                document,
                page,
                layout,
                regular_size=10,
                yellow_rectangles=[],
                border_rectangles=[],
            )
    finally:
        source.close()
    return document


def _make_geometry_pdf() -> bytes:
    source = fitz.open()
    page = source.new_page(width=600, height=800)
    justified_text = " ".join(
        [
            "Lorem ipsum dolor sit amet consectetur adipiscing elit sed do "
            "eiusmod tempor incididunt ut labore et dolore magna aliqua"
        ]
        * 3
    )
    page.insert_textbox(
        fitz.Rect(50, 90, 550, 210),
        justified_text,
        fontsize=10,
        fontname="helv",
        align=fitz.TEXT_ALIGN_JUSTIFY,
        lineheight=1.2,
    )
    for text, baseline in (
        ("LEFT GEOMETRY FIRST LINE", 270),
        ("LEFT GEOMETRY SECOND", 282),
        ("LEFT END", 294),
    ):
        page.insert_text(
            (50, baseline),
            text,
            fontsize=10,
            fontname="helv",
        )
    page.insert_textbox(
        fitz.Rect(100, 350, 500, 390),
        "CENTER GEOMETRY\nSHORT CENTER",
        fontsize=10,
        fontname="helv",
        align=fitz.TEXT_ALIGN_CENTER,
        lineheight=1.2,
    )
    page.insert_textbox(
        fitz.Rect(100, 430, 550, 470),
        "RIGHT GEOMETRY\nRIGHT END",
        fontsize=10,
        fontname="helv",
        align=fitz.TEXT_ALIGN_RIGHT,
        lineheight=1.2,
    )
    content = source.tobytes(garbage=4, deflate=True)
    source.close()
    return content


def test_alignment_geometry_maps_to_real_word_alignment() -> None:
    justified = _analyze(
        [
            _line("justified line one", left=50, right=550, baseline=100),
            _line("justified line two", left=50.5, right=549.5, baseline=112),
            _line("justified line three", left=50, right=550, baseline=124),
            _line("short last line", left=50, right=220, baseline=136),
        ]
    )[0]
    left = _analyze(
        [
            _line("left one", left=50, right=550, baseline=200),
            _line("left two", left=50, right=430, baseline=212),
            _line("left three", left=50, right=500, baseline=224),
            _line("left four", left=50, right=260, baseline=236),
        ]
    )[0]
    centered = _analyze(
        [
            _line("center one", left=200, right=400, baseline=300),
            _line("center two", left=170, right=430, baseline=312),
        ]
    )[0]
    right = _analyze(
        [
            _line("right one", left=300, right=550, baseline=380),
            _line("right two", left=350, right=550, baseline=392),
            _line("right three", left=280, right=550, baseline=404),
        ]
    )[0]

    assert justified.alignment == "justify"
    assert left.alignment == "left"
    assert centered.alignment == "center"
    assert right.alignment == "right"

    document = _write_layouts([justified, left, centered, right])

    assert [paragraph.alignment for paragraph in document.paragraphs] == [
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.RIGHT,
    ]


def test_pdf_fixture_produces_real_word_paragraph_alignments(
    tmp_path: Path,
) -> None:
    source = tmp_path / "paragraph-geometry.pdf"
    output = tmp_path / "paragraph-geometry.docx"
    source.write_bytes(_make_geometry_pdf())

    PdfToDocxConverter().convert(source, output)
    document = Document(output)
    paragraphs = {
        paragraph.text: paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    }
    justified = next(
        paragraph
        for text, paragraph in paragraphs.items()
        if text.startswith("Lorem ipsum")
    )
    left = next(
        paragraph
        for text, paragraph in paragraphs.items()
        if text.startswith("LEFT GEOMETRY")
    )
    centered = next(
        paragraph
        for text, paragraph in paragraphs.items()
        if text.startswith("CENTER GEOMETRY")
    )
    right = next(
        paragraph
        for text, paragraph in paragraphs.items()
        if text.startswith("RIGHT GEOMETRY")
    )

    assert justified.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert left.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert centered.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert right.alignment == WD_ALIGN_PARAGRAPH.RIGHT


@pytest.mark.parametrize(
    ("source_ratio", "expected_word_ratio"),
    [(1.0, 1.0), (1.15, 1.15), (1.5, 1.5), (2.0, 2.0)],
)
def test_baseline_spacing_is_normalized_to_natural_word_values(
    source_ratio: float,
    expected_word_ratio: float,
) -> None:
    baseline_step = 10 * source_ratio
    layout = _analyze(
        [
            _line("spacing one", left=50, right=550, baseline=100),
            _line(
                "spacing two",
                left=50,
                right=550,
                baseline=100 + baseline_step,
            ),
            _line(
                "spacing three",
                left=50,
                right=250,
                baseline=100 + baseline_step * 2,
            ),
        ]
    )[0]
    document = _write_layouts([layout])
    paragraph = document.paragraphs[0]

    assert layout.source_line_spacing_ratio == pytest.approx(source_ratio)
    assert paragraph.paragraph_format.line_spacing == expected_word_ratio
    assert not document.element.xpath('.//w:spacing[@w:lineRule="exact"]')


def test_micro_variations_use_the_median_and_same_spacing_category() -> None:
    layouts = [
        _analyze(
            [
                _line("one", left=50, right=550, baseline=100),
                _line("two", left=50, right=550, baseline=100 + first_step),
                _line(
                    "three",
                    left=50,
                    right=250,
                    baseline=100 + first_step + second_step,
                ),
            ]
        )[0]
        for first_step, second_step in ((11.2, 11.6), (11.5, 11.7), (11.7, 11.4))
    ]

    assert [layout.line_spacing_ratio for layout in layouts] == [1.15, 1.15, 1.15]
    assert statistics.median(
        layout.source_line_spacing_ratio for layout in layouts
    ) == pytest.approx(1.15, abs=0.03)


def test_left_right_first_line_and_hanging_indents_are_word_properties() -> None:
    no_indent = _analyze(
        [
            _line("plain one", left=50, right=550, baseline=100),
            _line("plain two", left=50, right=550, baseline=112),
            _line("plain end", left=50, right=260, baseline=124),
        ]
    )[0]
    left_indent = _analyze(
        [
            _line("left indent one", left=90, right=550, baseline=180),
            _line("left indent two", left=90, right=550, baseline=192),
            _line("left indent end", left=90, right=280, baseline=204),
        ]
    )[0]
    right_indent = _analyze(
        [
            _line("right indent one", left=50, right=500, baseline=260),
            _line("right indent two", left=50, right=500, baseline=272),
            _line("right indent end", left=50, right=260, baseline=284),
        ]
    )[0]
    first_line = _analyze(
        [
            _line("first line", left=80, right=550, baseline=340),
            _line("continuation", left=50, right=550, baseline=352),
            _line("ending", left=50, right=260, baseline=364),
        ]
    )[0]
    hanging = _analyze(
        [
            _line("Ref. opening", left=50, right=550, baseline=420),
            _line("continuation", left=80, right=550, baseline=432),
            _line("ending", left=80, right=280, baseline=444),
        ]
    )[0]

    assert no_indent.left_indent == 0
    assert no_indent.right_indent == 0
    assert left_indent.left_indent == 40
    assert right_indent.right_indent == 50
    assert first_line.first_line_indent == 30
    assert hanging.left_indent == 30
    assert hanging.first_line_indent == -30

    document = _write_layouts(
        [no_indent, left_indent, right_indent, first_line, hanging]
    )
    paragraphs = document.paragraphs
    assert paragraphs[1].paragraph_format.left_indent.pt == pytest.approx(40)
    assert paragraphs[2].paragraph_format.right_indent.pt == pytest.approx(50)
    assert paragraphs[3].paragraph_format.first_line_indent.pt == pytest.approx(30)
    assert paragraphs[4].paragraph_format.left_indent.pt == pytest.approx(30)
    assert paragraphs[4].paragraph_format.first_line_indent.pt == pytest.approx(-30)


def test_interparagraph_gap_becomes_spacing_not_line_spacing() -> None:
    layouts = _analyze(
        [
            _line("paragraph A1", left=50, right=550, baseline=100),
            _line("paragraph A2", left=50, right=550, baseline=112),
            _line("paragraph A3", left=50, right=260, baseline=124),
            _line("paragraph B1", left=50, right=550, baseline=160),
            _line("paragraph B2", left=50, right=260, baseline=172),
        ]
    )

    assert len(layouts) == 2
    assert layouts[0].space_after == 3
    assert layouts[0].line_spacing_ratio == 1.15
    assert layouts[1].line_spacing_ratio == 1.15
    document = _write_layouts(list(layouts))
    assert document.paragraphs[0].paragraph_format.space_after.pt == pytest.approx(3)
    assert document.paragraphs[0].paragraph_format.line_spacing == 1.15


def test_style_change_columns_and_lists_are_not_merged() -> None:
    layouts = _analyze(
        [
            _line(
                "A centered title",
                left=180,
                right=420,
                baseline=100,
                font_size=18,
                bold=True,
            ),
            _line("body line", left=50, right=300, baseline=122),
            _line("body continuation", left=50, right=300, baseline=134),
            _line("other column", left=350, right=550, baseline=134),
            _line("- list item", left=50, right=300, baseline=170),
            _line("list continuation", left=70, right=300, baseline=182),
        ]
    )

    assert len(layouts) == 4
    assert len(layouts[0].lines) == 1
    assert len(layouts[1].lines) == 2
    assert len(layouts[2].lines) == 1
    assert len(layouts[3].lines) == 2


def test_separate_list_marker_and_text_on_same_baseline_stay_together() -> None:
    layouts = _analyze(
        [
            _line("- ", left=50, right=58, baseline=100),
            _line("list item", left=70, right=300, baseline=100),
            _line("list continuation", left=70, right=300, baseline=112),
        ]
    )

    assert len(layouts) == 1
    assert len(layouts[0].lines) == 2
    assert len(layouts[0].lines[0].fragments) == 2
    assert layouts[0].lines[0].text == "- list item"


def test_ocr_edge_tolerance_keeps_noisy_justified_lines_together() -> None:
    lines = [
        _line("ocr one", left=50, right=550, baseline=100),
        _line("ocr two", left=53, right=546, baseline=112),
        _line("ocr three", left=49, right=551, baseline=124),
        _line("ocr ending", left=52, right=230, baseline=136),
    ]

    native = _analyze(lines, origin="native")[0]
    ocr = _analyze(lines, origin="ocr")[0]

    assert native.alignment == "left"
    assert ocr.alignment == "justify"


def test_reference_fixture_keeps_content_styles_and_natural_spacing(
    tmp_path: Path,
) -> None:
    source = (
        Path(__file__).resolve().parents[3]
        / "apps"
        / "web"
        / "e2e"
        / "fixtures"
        / "conversion-docx-fidelity.pdf"
    )
    output = tmp_path / "reference.docx"

    PdfToDocxConverter().convert(source, output)
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    ordinary_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.style.name == "Normal" and paragraph.text.strip()
    ]

    assert "Engagement individuel" in text
    assert ordinary_paragraphs
    assert all(
        paragraph.alignment in {
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        for paragraph in ordinary_paragraphs
    )
    assert any(
        paragraph.paragraph_format.left_indent is not None
        for paragraph in ordinary_paragraphs
    )
    assert all(
        isinstance(paragraph.paragraph_format.line_spacing, float)
        for paragraph in ordinary_paragraphs
    )
    assert not document.element.xpath('.//w:spacing[@w:lineRule="exact"]')


@pytest.mark.integration
@pytest.mark.skipif(
    not LIBREOFFICE_AVAILABLE,
    reason="LibreOffice est requis pour la validation visuelle DOCX.",
)
def test_rendered_paragraph_geometry_keeps_density_and_page_count(
    tmp_path: Path,
) -> None:
    source_path = tmp_path / "paragraph-geometry.pdf"
    docx_path = tmp_path / "paragraph-geometry-rendered.docx"
    source_path.write_bytes(_make_geometry_pdf())
    PdfToDocxConverter().convert(source_path, docx_path)
    profile = tmp_path / "lo-profile"
    runtime = tmp_path / "lo-runtime"
    runtime.mkdir(mode=0o700)

    completed = subprocess.run(
        [
            shutil.which("libreoffice") or "libreoffice",
            "--headless",
            f"-env:UserInstallation={profile.as_uri()}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(tmp_path),
            str(docx_path),
        ],
        check=False,
        capture_output=True,
        text=True,
        timeout=45,
        env={**os.environ, "XDG_RUNTIME_DIR": str(runtime)},
    )
    rendered_path = tmp_path / "paragraph-geometry-rendered.pdf"
    if completed.returncode != 0 or not rendered_path.is_file():
        pytest.skip("Le rendu LibreOffice headless n'est pas disponible.")

    with fitz.open(source_path) as source, fitz.open(rendered_path) as rendered:
        assert len(source) == len(rendered) == 1
        source_text = source[0].get_text("text")
        rendered_text = rendered[0].get_text("text")
        assert "Lorem ipsum" in rendered_text
        assert "CENTER GEOMETRY" in rendered_text
        assert len(rendered_text.split()) >= len(source_text.split()) * 0.95
        rendered_blocks = [
            block
            for block in rendered[0].get_text("blocks", sort=True)
            if str(block[4]).strip()
        ]
        assert max(block[3] for block in rendered_blocks) < rendered[0].rect.height * 0.9
