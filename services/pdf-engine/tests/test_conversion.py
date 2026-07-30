from __future__ import annotations

import asyncio
import hashlib
import json
import tempfile
import zipfile
from pathlib import Path
from typing import Any

import fitz
import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from fastapi import UploadFile
from starlette.datastructures import Headers

from app import ocr
from app.conversion.docx_converter import (
    DOCX_EDITABLE_DEGRADED_WARNING,
    PdfToDocxConverter,
    text_retention_ratio,
)
from app.conversion.errors import ConversionError
from app.conversion.image_converter import PdfToImageConverter
from app.conversion.models import (
    ConversionOptions,
    DocxMode,
    OcrMode,
    TargetFormat,
    parse_page_range,
)
from app.conversion.routes import convert_pdf
from app.conversion.service import (
    build_options,
    cleanup_temporary_directory,
    prepare_conversion,
)
from app.conversion.text_converter import (
    PdfToHtmlConverter,
    PdfToTextConverter,
)

REFERENCE_FIXTURE_DIR = (
    Path(__file__).resolve().parents[3] / "apps" / "web" / "e2e" / "fixtures"
)


def make_upload(content: bytes, name: str = "source.pdf") -> UploadFile:
    upload_content = tempfile.SpooledTemporaryFile()
    upload_content.write(content)
    upload_content.seek(0)
    return UploadFile(
        file=upload_content,
        filename=name,
        headers=Headers({"content-type": "application/pdf"}),
    )


def save_document(document: fitz.Document) -> bytes:
    content = document.tobytes(garbage=4, deflate=True)
    document.close()
    return content


def make_text_pdf(
    texts: list[str],
    *,
    landscape_page: int | None = None,
) -> bytes:
    document = fitz.open()
    for index, text in enumerate(texts):
        rectangle = (
            fitz.Rect(0, 0, 792, 612)
            if landscape_page == index
            else fitz.Rect(0, 0, 612, 792)
        )
        page = document.new_page(width=rectangle.width, height=rectangle.height)
        page.insert_textbox(
            fitz.Rect(60, 70, rectangle.width - 60, rectangle.height - 70),
            text,
            fontsize=18,
            fontname="helv",
        )
    return save_document(document)


def make_scan_pdf(text: str = "SCANNED WITNESS ENGLISH") -> bytes:
    source = fitz.open()
    source_page = source.new_page(width=612, height=792)
    source_page.insert_text(
        (55, 220),
        text,
        fontsize=30,
        fontname="helv",
    )
    image = source_page.get_pixmap(dpi=200, alpha=False).tobytes("png")
    source.close()

    scanned = fitz.open()
    scanned_page = scanned.new_page(width=612, height=792)
    scanned_page.insert_image(scanned_page.rect, stream=image)
    return save_document(scanned)


def make_mixed_pdf() -> bytes:
    document = fitz.open()
    digital_page = document.new_page(width=612, height=792)
    digital_page.insert_text(
        (60, 160),
        "DIGITAL WITNESS",
        fontsize=24,
        fontname="helv",
    )
    scan = fitz.open(stream=make_scan_pdf(), filetype="pdf")
    document.insert_pdf(scan)
    scan.close()
    return save_document(document)


def make_image_backed_text_pdf() -> bytes:
    background_source = fitz.open()
    background_page = background_source.new_page(width=612, height=792)
    background_page.draw_rect(
        background_page.rect,
        fill=(0.96, 0.96, 0.96),
    )
    background_page.insert_text(
        (60, 160),
        "BACKGROUND RENDERING",
        fontsize=24,
    )
    background = background_page.get_pixmap(dpi=96, alpha=False).tobytes("png")
    background_source.close()

    document = fitz.open()
    page = document.new_page(width=612, height=792)
    page.insert_image(page.rect, stream=background)
    page.insert_text(
        (60, 160),
        "EDITABLE TEXT LAYER",
        fontsize=24,
    )
    return save_document(document)


def make_table_and_image_pdf() -> bytes:
    image_source = fitz.open()
    image_page = image_source.new_page(width=160, height=100)
    image_page.draw_rect(image_page.rect, fill=(0.15, 0.45, 0.85))
    image_page.insert_text((20, 60), "IMAGE", fontsize=20, color=(1, 1, 1))
    image = image_page.get_pixmap(alpha=False).tobytes("png")
    image_source.close()

    document = fitz.open()
    page = document.new_page(width=612, height=792)
    page.insert_text((60, 60), "Rapport illustré", fontsize=24)
    x_positions = [60, 220, 380]
    y_positions = [140, 190, 240]
    for x_position in x_positions:
        page.draw_line((x_position, y_positions[0]), (x_position, y_positions[-1]))
    for y_position in y_positions:
        page.draw_line((x_positions[0], y_position), (x_positions[-1], y_position))
    page.insert_text((75, 175), "Nom", fontsize=12)
    page.insert_text((235, 175), "Valeur", fontsize=12)
    page.insert_text((75, 225), "Été", fontsize=12)
    page.insert_text((235, 225), "42", fontsize=12)
    page.insert_image(fitz.Rect(60, 300, 300, 450), stream=image)
    return save_document(document)


def write_fixture(tmp_path: Path, content: bytes) -> Path:
    path = tmp_path / "source.pdf"
    path.write_bytes(content)
    return path


def test_page_range_is_deterministic_and_rejects_invalid_values() -> None:
    assert parse_page_range(None, 5) == [0, 1, 2, 3, 4]
    assert parse_page_range("1-3,5,2", 5) == [0, 1, 2, 4]
    for value in ("0", "6", "3-1", "one", "1,,2"):
        with pytest.raises(ConversionError, match="plage"):
            parse_page_range(value, 5)


def test_invalid_target_and_dpi_have_stable_validation_errors() -> None:
    with pytest.raises(ConversionError) as target_error:
        build_options(
            target_format="xlsx",
            languages="fra",
            ocr_mode="auto",
            pages=None,
            image_dpi=150,
            image_quality=85,
        )
    assert target_error.value.code == "UNSUPPORTED_TARGET_FORMAT"

    with pytest.raises(ConversionError) as dpi_error:
        build_options(
            target_format="png",
            languages="fra",
            ocr_mode="auto",
            pages=None,
            image_dpi=600,
            image_quality=85,
        )
    assert dpi_error.value.code == "CONVERSION_FAILED"
    assert "96, 150 ou 300" in dpi_error.value.message


def test_txt_is_utf8_ordered_and_separates_pages(tmp_path: Path) -> None:
    source = write_fixture(
        tmp_path,
        make_text_pdf(["Été à Paris", "Second page in English"]),
    )
    output = tmp_path / "conversion.txt"

    artifact = PdfToTextConverter().convert(source, output)

    decoded = output.read_bytes().decode("utf-8")
    assert artifact.media_type.startswith("text/plain")
    assert "--- Page 1 ---" in decoded
    assert "--- Page 2 ---" in decoded
    assert "Été à Paris" in decoded
    assert decoded.index("Été") < decoded.index("Second page")


def test_html_is_autonomous_utf8_and_embeds_images(tmp_path: Path) -> None:
    source = write_fixture(tmp_path, make_table_and_image_pdf())
    output = tmp_path / "conversion.html"

    PdfToHtmlConverter().convert(source, output)

    decoded = output.read_text(encoding="utf-8")
    assert decoded.startswith("<!doctype html>")
    assert '<meta charset="utf-8">' in decoded
    assert "Rapport illustré" in decoded
    assert "data:image/png;base64," in decoded
    assert "http://" not in decoded
    assert "https://" not in decoded


def test_docx_is_openable_with_text_image_table_and_orientation(
    tmp_path: Path,
) -> None:
    table_source = write_fixture(tmp_path, make_table_and_image_pdf())
    table_output = tmp_path / "table.docx"
    PdfToDocxConverter().convert(table_source, table_output)
    table_document = Document(table_output)

    assert zipfile.is_zipfile(table_output)
    assert "Rapport illustré" in "\n".join(
        paragraph.text for paragraph in table_document.paragraphs
    )
    assert table_document.inline_shapes
    assert table_document.tables
    assert any(
        "Été" in cell.text
        for table in table_document.tables
        for row in table.rows
        for cell in row.cells
    )

    oriented_source = write_fixture(
        tmp_path,
        make_text_pdf(
            ["Portrait page", "Landscape page"],
            landscape_page=1,
        ),
    )
    oriented_output = tmp_path / "oriented.docx"
    PdfToDocxConverter().convert(oriented_source, oriented_output)
    oriented_document = Document(oriented_output)
    assert len(oriented_document.sections) == 2
    assert oriented_document.sections[0].page_height > (
        oriented_document.sections[0].page_width
    )
    assert oriented_document.sections[1].page_width > (
        oriented_document.sections[1].page_height
    )


@pytest.mark.parametrize(
    (
        "fixture_name",
        "expected_texts",
        "expected_images",
        "expected_tables",
        "expected_landscape",
    ),
    [
        (
            "conversion-simple-text.pdf",
            ("Conversion locale PDF",),
            0,
            0,
            False,
        ),
        (
            "conversion-scan.pdf",
            ("SCANNED OCR WITNESS",),
            0,
            0,
            False,
        ),
        (
            "conversion-mixed.pdf",
            ("DIGITAL MIXED WITNESS", "SCANNED MIXED WITNESS"),
            0,
            0,
            False,
        ),
        (
            "conversion-images.pdf",
            ("Document avec image intégrée",),
            1,
            0,
            False,
        ),
        (
            "conversion-table.pdf",
            ("Tableau simple",),
            0,
            1,
            False,
        ),
        (
            "conversion-landscape.pdf",
            ("Landscape conversion witness",),
            0,
            0,
            True,
        ),
    ],
)
def test_reference_docx_corpus_is_openable_and_contains_expected_content(
    fixture_name: str,
    expected_texts: tuple[str, ...],
    expected_images: int,
    expected_tables: int,
    expected_landscape: bool,
) -> None:
    source = REFERENCE_FIXTURE_DIR / fixture_name
    prepared = asyncio.run(
        prepare_conversion(
            make_upload(source.read_bytes(), fixture_name),
            ConversionOptions(
                target_format=TargetFormat.DOCX,
                languages="eng",
                ocr_mode=OcrMode.AUTO,
            ),
        )
    )
    try:
        assert prepared.artifact.path.stat().st_size > 0
        assert zipfile.is_zipfile(prepared.artifact.path)
        document = Document(prepared.artifact.path)
        extracted_text = "\n".join(
            [
                *(paragraph.text for paragraph in document.paragraphs),
                *(
                    cell.text
                    for table in document.tables
                    for row in table.rows
                    for cell in row.cells
                ),
            ]
        )
        for expected_text in expected_texts:
            assert expected_text in extracted_text
        assert len(document.inline_shapes) >= expected_images
        assert len(document.tables) >= expected_tables
        assert any(
            section.page_width > section.page_height
            for section in document.sections
        ) is expected_landscape
    finally:
        cleanup_temporary_directory(prepared.temporary_directory)


def test_editable_docx_preserves_fidelity_markers(tmp_path: Path) -> None:
    source = REFERENCE_FIXTURE_DIR / "conversion-docx-fidelity.pdf"
    output = tmp_path / "editable.docx"

    PdfToDocxConverter().convert(source, output, mode=DocxMode.EDITABLE)

    assert output.stat().st_size > 0
    assert zipfile.is_zipfile(output)
    document = Document(output)
    assert len(document.sections) == 3
    assert len(document.inline_shapes) == 3
    assert all(
        100 <= image.width.pt <= 140
        and 25 <= image.height.pt <= 50
        for image in document.inline_shapes
    )

    title = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "Engagement individuel"
    )
    title_index = next(
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text == "Engagement individuel"
    )
    assert title.style.name == "Title"
    assert title.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert any(run.bold for run in title.runs)
    subtitle = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip()
        == "Exemplaire à remettre signé à l'administration"
    )
    subtitle_index = next(
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip()
        == "Exemplaire à remettre signé à l'administration"
    )
    assert subtitle_index == title_index + 1
    assert subtitle.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert not any(
        paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
        and len(paragraph.text.split()) >= 12
        for paragraph in document.paragraphs
    )
    assert any(
        run.bold
        for paragraph in document.paragraphs
        if "respect de ces règles" in paragraph.text
        for run in paragraph.runs
    )
    assert any(
        run.font.highlight_color == WD_COLOR_INDEX.YELLOW
        for paragraph in document.paragraphs
        if "Nom de l'étudiant" in paragraph.text
        for run in paragraph.runs
    )
    assert any(
        paragraph._p.xpath(".//w:pBdr")
        for paragraph in document.paragraphs
        if "premier paragraphe encadré" in paragraph.text
    )
    assert sum(
        paragraph.style.name == "List Bullet"
        for paragraph in document.paragraphs
    ) >= 3
    assert not any(
        paragraph.style.name in {"List Bullet", "List Number"}
        and not paragraph.text.strip()
        for paragraph in document.paragraphs
    )
    first_drawing_index = next(
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph._p.xpath(".//w:drawing")
    )
    assert first_drawing_index < title_index
    assert isinstance(title.paragraph_format.line_spacing, float)
    assert title.paragraph_format.line_spacing >= 1.03
    assert title.paragraph_format.space_after is not None
    assert title.paragraph_format.space_after.pt > 0
    assert not document.element.xpath(
        './/w:spacing[@w:lineRule="exact"]'
    )

    long_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if len(paragraph.text.split()) >= 20
    ]
    assert long_paragraphs
    assert all(
        isinstance(paragraph.paragraph_format.line_spacing, float)
        and paragraph.paragraph_format.line_spacing >= 1.12
        for paragraph in long_paragraphs
    )
    list_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.style.name in {"List Bullet", "List Number"}
    ]
    assert all(
        isinstance(paragraph.paragraph_format.line_spacing, float)
        and paragraph.paragraph_format.line_spacing >= 1.05
        for paragraph in list_paragraphs
    )
    assert any(
        paragraph.paragraph_format.space_before is not None
        and paragraph.paragraph_format.space_before.pt > 0
        for paragraph in list_paragraphs
    )

    bordered_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if "premier paragraphe encadré" in paragraph.text
    )
    border_edges = bordered_paragraph._p.xpath("./w:pPr/w:pBdr/*")
    assert border_edges
    assert all(int(edge.get(qn("w:space"), "0")) >= 7 for edge in border_edges)
    assert bordered_paragraph.paragraph_format.space_before is not None
    assert bordered_paragraph.paragraph_format.space_before.pt >= 3
    assert bordered_paragraph.paragraph_format.space_after is not None
    assert bordered_paragraph.paragraph_format.space_after.pt >= 3

    with zipfile.ZipFile(output) as archive:
        logo = fitz.Pixmap(archive.read("word/media/image1.png"))
    assert logo.alpha
    color_channels = min(3, logo.n)
    logo_samples = logo.samples
    black_pixels = sum(
        (
            logo.n < 4 or logo_samples[index + logo.n - 1] >= 32
        )
        and all(
            logo_samples[index + channel] < 24
            for channel in range(color_channels)
        )
        for index in range(0, len(logo_samples), logo.n)
    )
    assert black_pixels / (logo.width * logo.height) < 0.05


def test_docx_image_soft_mask_drops_existing_alpha_before_composition() -> None:
    source = REFERENCE_FIXTURE_DIR / "conversion-docx-fidelity.pdf"
    with fitz.open(source) as pdf:
        page = pdf[0]
        block = next(
            item
            for item in page.get_text("dict")["blocks"]
            if item.get("type") == 1
        )
        color_pixmap = fitz.Pixmap(block["image"])
        color_with_alpha = fitz.Pixmap(color_pixmap, 1)
        block_with_redundant_alpha = {
            **block,
            "image": color_with_alpha.tobytes("png"),
        }

        rendered = PdfToDocxConverter._render_image(
            page,
            block_with_redundant_alpha,
            fitz.Rect(block["bbox"]),
        )

    result = fitz.Pixmap(rendered)
    assert result.colorspace is not None
    assert result.colorspace.n == 3
    assert result.alpha
    assert result.width == color_pixmap.width
    assert result.height == color_pixmap.height


def test_docx_runs_keep_mixed_bold_formatting() -> None:
    document = Document()
    paragraph = document.add_paragraph()

    PdfToDocxConverter._append_styled_run(
        paragraph,
        {
            "text": "Texte normal ",
            "font": "ArialMT",
            "flags": 0,
            "size": 9,
            "color": 0,
        },
        "Texte normal ",
        highlighted=False,
    )
    PdfToDocxConverter._append_styled_run(
        paragraph,
        {
            "text": "expression importante",
            "font": "Arial-BoldMT",
            "flags": 16,
            "size": 9,
            "color": 0,
        },
        "expression importante",
        highlighted=False,
    )

    assert paragraph.runs[0].bold is False
    assert paragraph.runs[1].bold is True


def test_docx_image_extraction_falls_back_to_a_white_page_clip() -> None:
    source = REFERENCE_FIXTURE_DIR / "conversion-docx-fidelity.pdf"
    with fitz.open(source) as pdf:
        page = pdf[0]
        block = next(
            item
            for item in page.get_text("dict")["blocks"]
            if item.get("type") == 1
        )
        rendered = PdfToDocxConverter._render_image(
            page,
            {**block, "image": b"not-an-image", "mask": b"not-a-mask"},
            fitz.Rect(block["bbox"]),
        )

    result = fitz.Pixmap(rendered)
    assert result.alpha == 0
    assert result.colorspace is not None
    assert result.colorspace.n == 3
    assert result.width > 0
    assert result.height > 0


def test_visual_docx_renders_one_full_page_image_per_source_page(
) -> None:
    source = REFERENCE_FIXTURE_DIR / "conversion-docx-fidelity.pdf"
    prepared = asyncio.run(
        prepare_conversion(
            make_upload(source.read_bytes(), source.name),
            ConversionOptions(
                target_format=TargetFormat.DOCX,
                ocr_mode=OcrMode.AUTO,
                docx_mode=DocxMode.VISUAL,
            ),
        )
    )
    temporary_directory = prepared.temporary_directory
    try:
        output = prepared.artifact.path
        document = Document(output)
        assert zipfile.is_zipfile(output)
        assert output.stat().st_size > 0
        assert len(document.sections) == 3
        assert len(document.inline_shapes) == 3
        for image_index, image in enumerate(document.inline_shapes):
            section = document.sections[image_index]
            assert image.width.pt > 580 and image.height.pt > 820
            assert 0.97 <= image.width / section.page_width <= 1
            assert 0.97 <= image.height / section.page_height <= 1
            extent = image._inline.xpath("./wp:extent")[0]
            assert int(extent.get("cx")) == image.width
            assert int(extent.get("cy")) == image.height

        image_paragraphs = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph._p.xpath(".//w:drawing")
        ]
        assert len(image_paragraphs) == 3
        for paragraph in image_paragraphs:
            spacing = paragraph._p.xpath("./w:pPr/w:spacing")
            assert not spacing or spacing[0].get(qn("w:lineRule")) != "exact"

        media_parts = [
            part
            for part in document.part.package.parts
            if part.content_type.startswith("image/")
            and "thumbnail" not in str(part.partname)
        ]
        assert len(media_parts) == 3
        for part in media_parts:
            pixmap = fitz.Pixmap(part.blob)
            samples = pixmap.samples
            non_white_pixels = sum(
                any(
                    samples[index + channel] < 248
                    for channel in range(min(3, pixmap.n))
                )
                for index in range(0, len(samples), pixmap.n)
            )
            assert non_white_pixels / (pixmap.width * pixmap.height) > 0.002
        assert any(
            "moins facilement modifiable" in warning
            for warning in prepared.artifact.warnings
        )
    finally:
        cleanup_temporary_directory(temporary_directory)
    assert not temporary_directory.exists()


def test_docx_mode_validation_is_stable() -> None:
    options = build_options(
        target_format="docx",
        languages="fra",
        ocr_mode="auto",
        pages=None,
        image_dpi=150,
        image_quality=85,
        docx_mode="visual",
    )
    assert options.docx_mode == DocxMode.VISUAL

    with pytest.raises(ConversionError) as error:
        build_options(
            target_format="docx",
            languages="fra",
            ocr_mode="auto",
            pages=None,
            image_dpi=150,
            image_quality=85,
            docx_mode="../../visual",
        )
    assert error.value.code == "CONVERSION_FAILED"
    assert "editable ou visual" in error.value.message


def test_editable_text_retention_uses_source_token_coverage() -> None:
    source = "Titre important texte répété répété et contenu final"
    assert text_retention_ratio(source, source) == 1
    assert text_retention_ratio(source, "Titre important texte") < 0.5
    assert text_retention_ratio("", "") == 1


def test_editable_docx_warns_instead_of_falling_back_to_visual(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    source = write_fixture(
        tmp_path,
        make_text_pdf(
            [
                " ".join(
                    f"sourceword{index}" for index in range(80)
                )
            ]
        ),
    )
    output = tmp_path / "degraded.docx"

    def append_incomplete_text(
        _converter: PdfToDocxConverter,
        document: Any,
        _pdf: fitz.Document,
    ) -> None:
        document.add_paragraph("sourceword0")

    monkeypatch.setattr(
        PdfToDocxConverter,
        "_append_editable_pages",
        append_incomplete_text,
    )
    artifact = PdfToDocxConverter().convert(
        source,
        output,
        mode=DocxMode.EDITABLE,
    )
    document = Document(output)

    assert DOCX_EDITABLE_DEGRADED_WARNING in artifact.warnings
    assert len(document.inline_shapes) == 0
    assert document.paragraphs[-1].text == "sourceword0"


def test_editable_docx_does_not_embed_a_full_page_screenshot_when_text_exists(
    tmp_path: Path,
) -> None:
    source = write_fixture(tmp_path, make_image_backed_text_pdf())
    output = tmp_path / "editable-with-background.docx"

    PdfToDocxConverter().convert(source, output, mode=DocxMode.EDITABLE)
    document = Document(output)
    extracted_text = "\n".join(
        paragraph.text for paragraph in document.paragraphs
    )

    assert "EDITABLE TEXT LAYER" in extracted_text
    assert len(document.inline_shapes) == 0


@pytest.mark.parametrize(
    ("target_format", "extension"),
    [(TargetFormat.PNG, "png"), (TargetFormat.JPEG, "jpg")],
)
def test_images_have_expected_dimensions_and_deterministic_names(
    tmp_path: Path,
    target_format: TargetFormat,
    extension: str,
) -> None:
    source = write_fixture(
        tmp_path,
        make_text_pdf(["Page one", "Page two"], landscape_page=1),
    )
    output_directory = tmp_path / target_format.value
    output_directory.mkdir()

    artifact = PdfToImageConverter().convert(
        source,
        output_directory,
        target_format=target_format,
        dpi=96,
        quality=80,
    )

    assert artifact.media_type == "application/zip"
    with zipfile.ZipFile(artifact.path) as archive:
        assert archive.namelist() == [
            f"document_page_0001.{extension}",
            f"document_page_0002.{extension}",
        ]
        first = fitz.Pixmap(archive.read(archive.namelist()[0]))
        second = fitz.Pixmap(archive.read(archive.namelist()[1]))
        assert (first.width, first.height) == (816, 1056)
        assert (second.width, second.height) == (1056, 816)


def test_single_image_is_returned_without_zip(tmp_path: Path) -> None:
    source = write_fixture(tmp_path, make_text_pdf(["Single page"]))
    output_directory = tmp_path / "single"
    output_directory.mkdir()

    artifact = PdfToImageConverter().convert(
        source,
        output_directory,
        target_format=TargetFormat.PNG,
        dpi=150,
        quality=85,
    )

    assert artifact.filename == "conversion_page_0001.png"
    assert artifact.media_type == "image/png"
    assert fitz.Pixmap(artifact.path).width == 1275


@pytest.mark.parametrize(
    ("source_bytes", "target_format", "expected_text", "classification"),
    [
        (make_scan_pdf(), TargetFormat.DOCX, "SCANNED", "scanned"),
        (make_mixed_pdf(), TargetFormat.TXT, "DIGITAL", "mixed"),
    ],
)
def test_auto_ocr_converts_a_scan_and_mixed_pdf(
    source_bytes: bytes,
    target_format: TargetFormat,
    expected_text: str,
    classification: str,
) -> None:
    prepared = asyncio.run(
        prepare_conversion(
            make_upload(source_bytes),
            ConversionOptions(
                target_format=target_format,
                languages="eng",
                ocr_mode=OcrMode.AUTO,
            ),
        )
    )
    try:
        if target_format == TargetFormat.DOCX:
            converted_text = "\n".join(
                paragraph.text
                for paragraph in Document(prepared.artifact.path).paragraphs
            )
        else:
            converted_text = prepared.artifact.path.read_text(encoding="utf-8")
        assert prepared.result.ocr_used is True
        assert expected_text in converted_text.upper()
        assert prepared.result.text_layer.classification == classification
    finally:
        cleanup_temporary_directory(prepared.temporary_directory)


def test_never_ocr_warns_and_does_not_modify_source(tmp_path: Path) -> None:
    source = make_scan_pdf()
    source_digest = hashlib.sha256(source).hexdigest()
    prepared = asyncio.run(
        prepare_conversion(
            make_upload(source),
            ConversionOptions(
                target_format=TargetFormat.TXT,
                ocr_mode=OcrMode.NEVER,
            ),
        )
    )
    try:
        assert prepared.result.ocr_used is False
        assert any("sans OCR" in warning for warning in prepared.result.warnings)
        assert hashlib.sha256(source).hexdigest() == source_digest
    finally:
        cleanup_temporary_directory(prepared.temporary_directory)


def test_endpoint_returns_headers_and_cleans_temporary_files() -> None:
    response = asyncio.run(
        convert_pdf(
            file=make_upload(make_text_pdf(["Endpoint witness"])),
            target_format="txt",
            languages="eng",
            ocr_mode="auto",
        )
    )
    output_path = Path(response.path)

    assert response.media_type == "text/plain; charset=utf-8"
    assert response.headers["x-conversion-format"] == "txt"
    assert response.headers["x-conversion-ocr-used"] == "false"
    assert response.headers["x-conversion-stage"] == "completed"
    assert output_path.read_text(encoding="utf-8").endswith("\n")
    assert response.background is not None
    cleanup_temporary_directory(output_path.parent)
    assert not output_path.parent.exists()


def test_invalid_pdf_has_stable_error_and_is_cleaned(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    temporary_directory = tmp_path / "conversion-temp"

    def create_directory() -> Path:
        temporary_directory.mkdir()
        return temporary_directory

    monkeypatch.setattr(
        "app.conversion.service.create_temporary_directory",
        create_directory,
    )
    with pytest.raises(ConversionError) as error:
        asyncio.run(
            prepare_conversion(
                make_upload(b"not a pdf"),
                ConversionOptions(target_format=TargetFormat.TXT),
            )
        )
    assert error.value.code == "INVALID_PDF"
    assert error.value.stage == "pdf_validation"
    assert not temporary_directory.exists()


@pytest.mark.parametrize(
    ("failure", "code"),
    [
        (TimeoutError(), "CONVERSION_TIMEOUT"),
        (FileNotFoundError(), "DEPENDENCY_UNAVAILABLE"),
    ],
)
def test_worker_timeout_and_dependency_errors_are_stable(
    failure: Exception,
    code: str,
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    temporary_directory = tmp_path / code.lower()

    def create_directory() -> Path:
        temporary_directory.mkdir()
        return temporary_directory

    async def fail_capture(
        command: list[str],
        *,
        timeout_seconds: int,
    ) -> tuple[int, bytes, bytes]:
        assert command
        assert timeout_seconds > 0
        raise failure

    monkeypatch.setattr(
        "app.conversion.service.create_temporary_directory",
        create_directory,
    )
    monkeypatch.setattr(ocr, "capture_process", fail_capture)
    with pytest.raises(ConversionError) as error:
        asyncio.run(
            prepare_conversion(
                make_upload(make_text_pdf(["Failure witness"])),
                ConversionOptions(target_format=TargetFormat.TXT),
            )
        )
    assert error.value.code == code
    assert error.value.stage == "txt_generation"
    assert not temporary_directory.exists()


def test_output_size_limit_is_stable_and_cleans_temporary_files(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    temporary_directory = tmp_path / "output-limit"

    def create_directory() -> Path:
        temporary_directory.mkdir()
        return temporary_directory

    monkeypatch.setattr(
        "app.conversion.service.create_temporary_directory",
        create_directory,
    )
    monkeypatch.setattr("app.conversion.service.MAX_OUTPUT_SIZE_BYTES", 1)

    with pytest.raises(ConversionError) as error:
        asyncio.run(
            prepare_conversion(
                make_upload(make_text_pdf(["Output size witness"])),
                ConversionOptions(target_format=TargetFormat.TXT),
            )
        )
    assert error.value.code == "OUTPUT_TOO_LARGE"
    assert not temporary_directory.exists()


def test_malicious_upload_name_is_ignored(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    temporary_directory = tmp_path / "safe-conversion"

    def create_directory() -> Path:
        temporary_directory.mkdir()
        return temporary_directory

    monkeypatch.setattr(
        "app.conversion.service.create_temporary_directory",
        create_directory,
    )
    prepared = asyncio.run(
        prepare_conversion(
            make_upload(
                make_text_pdf(["Safe filename witness"]),
                "../../private/secret.pdf",
            ),
            ConversionOptions(target_format=TargetFormat.TXT),
        )
    )
    try:
        assert prepared.artifact.filename == "conversion.txt"
        assert prepared.artifact.path.parent == temporary_directory
        assert (temporary_directory / "input.pdf").is_file()
        assert not (tmp_path / "private").exists()
    finally:
        cleanup_temporary_directory(prepared.temporary_directory)


def test_worker_failure_is_sanitized(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    async def fail_capture(
        command: list[str],
        *,
        timeout_seconds: int,
    ) -> tuple[int, bytes, bytes]:
        temporary_directory = str(Path(command[command.index("--input") + 1]).parent)
        diagnostic = json.dumps({"temporary": temporary_directory}).encode()
        return 1, b"", diagnostic

    monkeypatch.setattr(ocr, "capture_process", fail_capture)
    with pytest.raises(ConversionError) as error:
        asyncio.run(
            prepare_conversion(
                make_upload(make_text_pdf(["Failure witness"])),
                ConversionOptions(target_format=TargetFormat.TXT),
            )
        )
    assert error.value.code == "CONVERSION_FAILED"
    assert error.value.stage == "txt_generation"
    assert "pdf-engine-conversion-" not in (error.value.diagnostic or "")


def test_unexpected_conversion_failure_reports_exact_stage_and_cleans_up(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    temporary_directory = tmp_path / "unexpected-failure"

    def create_directory() -> Path:
        temporary_directory.mkdir()
        return temporary_directory

    def fail_detection(*_args: object, **_kwargs: object) -> object:
        raise RuntimeError("synthetic detector failure")

    monkeypatch.setattr(
        "app.conversion.service.create_temporary_directory",
        create_directory,
    )
    monkeypatch.setattr(
        "app.conversion.service.detect_text_layer",
        fail_detection,
    )

    with pytest.raises(ConversionError) as error:
        asyncio.run(
            prepare_conversion(
                make_upload(make_text_pdf(["Stage witness"])),
                ConversionOptions(target_format=TargetFormat.DOCX),
            )
        )

    assert error.value.code == "CONVERSION_FAILED"
    assert error.value.status_code == 502
    assert error.value.stage == "text_layer_detection"
    assert error.value.diagnostic == "RuntimeError"
    assert not temporary_directory.exists()
