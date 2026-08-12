from __future__ import annotations

import os
import shutil
import subprocess
from pathlib import Path
from typing import Any

import fitz
import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.conversion.docx_converter import PdfToDocxConverter
from app.conversion.logical_lines import canonicalize_visual_lines
from app.conversion.paragraph_layout import analyze_paragraph_layouts


PAGE = fitz.Rect(0, 0, 600, 800)
CONTENT = fitz.Rect(72, 60, 530, 740)
LIBREOFFICE_AVAILABLE = shutil.which("libreoffice") is not None


def _fragment(
    text: str,
    *,
    left: float,
    right: float,
    baseline: float,
    font_size: float = 10,
    italic: bool = False,
    block: int = 0,
) -> dict[str, Any]:
    bbox = (left, baseline - font_size, right, baseline + font_size * 0.2)
    return {
        "bbox": bbox,
        "_source_block_index": block,
        "spans": [
            {
                "text": text,
                "bbox": bbox,
                "origin": (left, baseline),
                "font": "Helvetica-Oblique" if italic else "Helvetica",
                "size": font_size,
                "flags": 2 if italic else 0,
                "color": 0,
            }
        ],
    }


def test_same_baseline_fragments_become_one_ordered_logical_line() -> None:
    result = canonicalize_visual_lines(
        [
            _fragment("word3", left=190, right=230, baseline=257.44, block=2),
            _fragment("word1", left=72, right=105, baseline=257.44, block=0),
            _fragment("word4", left=240, right=280, baseline=257.44, block=3),
            _fragment("word2", left=115, right=180, baseline=257.44, block=1),
        ],
        PAGE.width,
    )

    assert len(result.lines) == 1
    assert result.same_baseline_fragments_merged == 3
    assert result.lines[0].text == "word1 word2 word3 word4"
    assert result.lines[0].source_block_indexes == (0, 1, 2, 3)


def test_same_baseline_keeps_styled_runs_instead_of_flattening_them() -> None:
    result = canonicalize_visual_lines(
        [
            _fragment("normal", left=72, right=115, baseline=100),
            _fragment(
                "(63% des répondants)",
                left=125,
                right=225,
                baseline=100.35,
                italic=True,
            ),
            _fragment("suite", left=235, right=270, baseline=100.1),
        ],
        PAGE.width,
    )

    line = result.lines[0]
    assert line.text == "normal (63% des répondants) suite"
    assert len(line.spans) == 3
    assert [int(span["flags"]) for span in line.spans] == [0, 2, 0]

    layouts = analyze_paragraph_layouts(result.lines, PAGE, CONTENT)
    document = Document()
    converter = PdfToDocxConverter()
    converter._configure_styles(document)
    source = fitz.open()
    page = source.new_page(width=PAGE.width, height=PAGE.height)
    try:
        converter._append_text_paragraph(
            document,
            page,
            layouts[0],
            regular_size=10,
            yellow_rectangles=[],
            border_rectangles=[],
        )
    finally:
        source.close()
    text_runs = [run for run in document.paragraphs[0].runs if run.text]
    assert [run.italic for run in text_runs] == [False, True, False]


def test_distinct_baselines_and_distant_columns_are_not_merged() -> None:
    different_rows = canonicalize_visual_lines(
        [
            _fragment("row one", left=72, right=180, baseline=100),
            _fragment("row two", left=72, right=180, baseline=120),
        ],
        PAGE.width,
    )
    distant_columns = canonicalize_visual_lines(
        [
            _fragment("left", left=50, right=150, baseline=200),
            _fragment("right", left=400, right=520, baseline=200),
        ],
        PAGE.width,
    )

    assert len(different_rows.lines) == 2
    assert len(distant_columns.lines) == 2


def test_ocr_uses_the_same_canonicalizer_with_wider_baseline_tolerance() -> None:
    fragments = [
        _fragment("ocr", left=72, right=100, baseline=100),
        _fragment("noise", left=110, right=150, baseline=101.05),
    ]

    native = canonicalize_visual_lines(fragments, PAGE.width, origin="native")
    ocr = canonicalize_visual_lines(fragments, PAGE.width, origin="ocr")

    assert len(native.lines) == 2
    assert len(ocr.lines) == 1


def test_cross_block_fragments_feed_one_indented_justified_paragraph() -> None:
    fragments: list[dict[str, Any]] = [
        _fragment(
            "Suite à cette enquête, plusieurs profils ont été",
            left=108,
            right=530,
            baseline=180,
            block=0,
        ),
        _fragment(
            "retenus. Ces entrepreneurs réalisent différents",
            left=72,
            right=530,
            baseline=192,
            block=1,
        ),
        _fragment(
            "travaux agricoles auprès de leurs clients et sur",
            left=72,
            right=530,
            baseline=204,
            block=2,
        ),
        _fragment(
            "une exploitation agricole. Ils peuvent assurer",
            left=72,
            right=530,
            baseline=216,
            block=3,
        ),
    ]
    words = "la prestation de service en application de produits phytosanitaires. Les".split()
    left = 72.0
    for index, word in enumerate(words):
        width = max(12, len(word) * 5.2)
        right = 530 if index == len(words) - 1 else left + width
        fragments.append(
            _fragment(
                word,
                left=left,
                right=right,
                baseline=228,
                block=4 + index,
            )
        )
        left += width + 3
    fragments.extend(
        [
            _fragment(
                "entrepreneurs ayant pour clientèle les exploitations ont",
                left=72,
                right=530,
                baseline=240,
                block=20,
            ),
            _fragment(
                "été exclus lorsque leur activité ne correspondait pas à",
                left=72,
                right=530,
                baseline=252,
                block=21,
            ),
            _fragment(
                "cette enquête.",
                left=72,
                right=155,
                baseline=264,
                block=22,
            ),
        ]
    )

    canonicalized = canonicalize_visual_lines(fragments, PAGE.width)
    layouts = analyze_paragraph_layouts(
        canonicalized.lines,
        PAGE,
        CONTENT,
    )

    assert canonicalized.same_baseline_fragments_merged == len(words) - 1
    assert len(canonicalized.lines) == 8
    assert len(layouts) == 1
    assert len(layouts[0].lines) == 8
    assert layouts[0].first_line_indent == 36
    assert layouts[0].alignment == "justify"
    assert "la prestation de service" in layouts[0].lines[4].text


def test_separate_list_marker_uses_general_same_baseline_rule() -> None:
    result = canonicalize_visual_lines(
        [
            _fragment("- ", left=60, right=66, baseline=300, block=0),
            _fragment(
                "le Grand-Est",
                left=78,
                right=145,
                baseline=300,
                block=1,
            ),
        ],
        PAGE.width,
    )

    assert len(result.lines) == 1
    assert result.lines[0].text == "- le Grand-Est"
    assert result.same_baseline_fragments_merged == 1


def _make_three_page_flow_pdf() -> bytes:
    def padded_line(text: str, start: float, target: float = 530) -> str:
        padding_words = (
            " dans",
            " le",
            " contexte",
            " du",
            " secteur",
            " agricole",
            " national",
            " élargi",
        )
        while True:
            remaining = target - start - fitz.get_text_length(
                text,
                fontname="helv",
                fontsize=10,
            )
            fitting_words = [
                word
                for word in padding_words
                if fitz.get_text_length(word, fontname="helv", fontsize=10)
                <= remaining
            ]
            if not fitting_words:
                return text
            text += max(
                fitting_words,
                key=lambda word: fitz.get_text_length(
                    word,
                    fontname="helv",
                    fontsize=10,
                ),
            )

    source = fitz.open()
    for page_number in range(1, 4):
        page = source.new_page(width=595, height=842)
        page.insert_text((72, 42), "FLOW REGRESSION", fontsize=9)
        page.insert_text((290, 815), str(page_number), fontsize=9)
        if page_number == 1:
            page.insert_text((72, 100), "FIRST PAGE CONTENT", fontsize=12)
        elif page_number == 3:
            page.insert_text((72, 100), "FINAL PAGE CONTENT", fontsize=12)
    page = source[1]
    page.insert_text((72, 95), "a. Répartition selon les régions", fontsize=13)
    paragraph_lines = (
        (108, 145, "Suite à cette enquête, plusieurs profils ont été retenus."),
        (72, 157, "Ces entrepreneurs réalisent différents travaux agricoles."),
        (72, 169, "Ils interviennent auprès de leurs clients et sur une exploitation."),
        (72, 181, "Ils peuvent assurer différents travaux et prestations."),
    )
    for x, baseline, text in paragraph_lines:
        page.insert_text((x, baseline), padded_line(text, x), fontsize=10)
    x = 72.0
    for word in (
        "la prestation de service en application de produits "
        "phytosanitaires. Les"
    ).split():
        page.insert_text((x, 193), word, fontsize=10)
        x += fitz.get_text_length(word, fontname="helv", fontsize=10) + 8
    page.insert_text((x, 193), padded_line("suite", x), fontsize=10)
    for baseline, text in (
        (205, "entrepreneurs ayant pour clientèle les exploitations ont été retenus."),
        (217, "Les profils hors périmètre ont été exclus de cette étude."),
    ):
        page.insert_text((72, baseline), padded_line(text, 72), fontsize=10)
    page.insert_text(
        (72, 229),
        "Cette organisation termine cette enquête.",
        fontsize=10,
    )
    for index, region in enumerate(
        ("le Grand-Est", "la Bretagne", "la Normandie", "l’Occitanie")
    ):
        baseline = 280 + index * 18
        page.insert_text((72, baseline), "- ", fontsize=10)
        page.insert_text((90, baseline), region, fontsize=10)
    graph = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 240, 100), False)
    graph.clear_with(210)
    page.insert_image(fitz.Rect(120, 390, 480, 540), stream=graph.tobytes("png"))

    content = source.tobytes(garbage=4, deflate=True)
    source.close()
    return content


@pytest.mark.integration
@pytest.mark.skipif(
    not LIBREOFFICE_AVAILABLE,
    reason="LibreOffice est requis pour valider la pagination DOCX.",
)
def test_same_baseline_flow_renders_three_source_pages_as_three_word_pages(
    tmp_path: Path,
) -> None:
    source_path = tmp_path / "flow-source.pdf"
    docx_path = tmp_path / "flow.docx"
    source_path.write_bytes(_make_three_page_flow_pdf())
    converter = PdfToDocxConverter()
    converter.convert(source_path, docx_path)
    document = Document(docx_path)

    matching_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if "la prestation de service en application" in paragraph.text
    ]
    assert len(matching_paragraphs) == 1
    assert matching_paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert matching_paragraphs[0].paragraph_format.first_line_indent is not None
    assert matching_paragraphs[0].paragraph_format.first_line_indent.pt == pytest.approx(36)
    assert converter.last_flow_metrics["same_baseline_fragments_merged"] == 14
    assert sum(
        paragraph.style.name == "List Bullet"
        for paragraph in document.paragraphs
    ) == 4

    profile = tmp_path / "lo-profile"
    runtime = tmp_path / "lo-runtime"
    runtime.mkdir(mode=0o700)
    completed = subprocess.run(
        [
            shutil.which("libreoffice") or "libreoffice",
            "--headless",
            f"-env:UserInstallation={profile.as_uri()}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(tmp_path),
            str(docx_path),
        ],
        check=False,
        capture_output=True,
        text=True,
        timeout=45,
        env={**os.environ, "XDG_RUNTIME_DIR": str(runtime)},
    )
    rendered_path = tmp_path / "flow.pdf"
    if completed.returncode != 0 or not rendered_path.is_file():
        pytest.skip("Le rendu LibreOffice headless n'est pas disponible.")

    with fitz.open(source_path) as source, fitz.open(rendered_path) as rendered:
        assert len(source) == len(rendered) == 3
        page_two_text = rendered[1].get_text("text")
        normalized_page_two_text = " ".join(page_two_text.split())
        assert "la prestation de service en application" in normalized_page_two_text
        assert "a. Répartition selon les régions" in page_two_text
        assert all(
            region in page_two_text
            for region in ("Grand-Est", "Bretagne", "Normandie", "Occitanie")
        )
        assert rendered[1].get_images(full=True)
        assert "FINAL PAGE CONTENT" in rendered[2].get_text("text")
