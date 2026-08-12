from __future__ import annotations

import io
import logging
import re
import statistics
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Sequence

import fitz
from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.conversion.header_footer import (
    ImageLayoutBlock,
    LayoutBlock,
    PageLayout,
    PageLayoutInput,
    PageNumberCandidate,
    TextOrigin,
    detect_document_layout,
)
from app.conversion.models import ConversionArtifact, DocxMode
from app.conversion.logical_lines import LogicalLine, canonicalize_visual_lines
from app.conversion.paragraph_layout import (
    LIST_MARKER,
    ParagraphLayout,
    analyze_paragraph_layouts,
    is_list_line,
    normalize_paragraph_gap,
)


LOGGER = logging.getLogger(__name__)

DOCX_LAYOUT_WARNING = (
    "La conversion tente de conserver la mise en page, mais certains éléments "
    "complexes peuvent être réorganisés."
)
DOCX_VISUAL_WARNING = (
    "Le mode fidèle visuellement conserve chaque page comme une image ; "
    "le contenu est moins facilement modifiable."
)
DOCX_EDITABLE_DEGRADED_WARNING = (
    "La conversion Word éditable est dégradée : moins de 50 % du texte "
    "source a été reconstruit. Utilisez le PDF source pour vérifier le contenu."
)
@dataclass(frozen=True)
class CoverPageAnalysis:
    is_cover_page: bool
    text_block_count: int
    word_count: int
    whitespace_ratio: float
    largest_font_size: float
    largest_vertical_gap: float
    title_position_ratio: float
    has_lower_page_text: bool


def text_retention_ratio(source_text: str, converted_text: str) -> float:
    source_tokens = Counter(re.findall(r"\w+", source_text.casefold()))
    if not source_tokens:
        return 1.0
    converted_tokens = Counter(re.findall(r"\w+", converted_text.casefold()))
    retained_tokens = source_tokens & converted_tokens
    return sum(retained_tokens.values()) / sum(source_tokens.values())


class PdfToDocxConverter:
    def convert(
        self,
        input_pdf: Path,
        output_docx: Path,
        *,
        mode: DocxMode = DocxMode.EDITABLE,
        text_origin: TextOrigin = "native",
    ) -> ConversionArtifact:
        self.last_flow_metrics = {
            "raw_text_fragments": 0,
            "logical_lines": 0,
            "same_baseline_fragments_merged": 0,
            "paragraphs": 0,
        }
        document = Document()
        with fitz.open(input_pdf) as pdf:
            if mode == DocxMode.VISUAL:
                self._append_visual_pages(document, pdf)
                warnings = (DOCX_VISUAL_WARNING,)
            else:
                source_text = "\n".join(
                    " ".join(
                        str(word[4])
                        for word in page.get_text("words", sort=True)
                    )
                    for page in pdf
                )
                self._configure_styles(document)
                self._text_origin = text_origin
                self._append_editable_pages(document, pdf)
                converted_text = self._document_text(document)
                retention_ratio = text_retention_ratio(
                    source_text,
                    converted_text,
                )
                warnings = (
                    (DOCX_LAYOUT_WARNING, DOCX_EDITABLE_DEGRADED_WARNING)
                    if retention_ratio < 0.5
                    else (DOCX_LAYOUT_WARNING,)
                )

        document.save(output_docx)
        return ConversionArtifact(
            path=output_docx,
            filename="conversion.docx",
            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            warnings=warnings,
        )

    @staticmethod
    def _document_text(document: DocumentType) -> str:
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        tables = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        stories = [
            element.text or ""
            for section in document.sections
            for story in (section.header, section.footer)
            for element in story._element.xpath(".//w:t")
        ]
        return "\n".join([*paragraphs, *tables, *stories])

    def _append_editable_pages(
        self,
        document: DocumentType,
        pdf: fitz.Document,
    ) -> None:
        text_origin = getattr(self, "_text_origin", "native")
        pages = tuple(pdf)
        page_blocks = tuple(
            page.get_text("dict", sort=True).get("blocks", []) for page in pages
        )
        layout = detect_document_layout(
            tuple(
                PageLayoutInput(
                    page_index=page_index,
                    width=page.rect.width,
                    height=page.rect.height,
                    blocks=blocks,
                    origin=text_origin,
                )
                for page_index, (page, blocks) in enumerate(
                    zip(pages, page_blocks, strict=True)
                )
            )
        )
        # A section is created exactly once at its final body boundary. Its
        # stories are then populated before that page's body; no later step
        # replaces its sectPr, headerReference, or footerReference.
        for page_index, (page, blocks, page_layout) in enumerate(
            zip(pages, page_blocks, layout.pages, strict=True)
        ):
            section = (
                document.sections[0]
                if page_index == 0
                else document.add_section(WD_SECTION.NEW_PAGE)
            )
            self._configure_editable_section(
                section,
                page,
                blocks,
                page_layout,
            )
            self._configure_header_footer(
                section,
                page,
                page_layout,
            )
            self._append_page_content(
                document,
                page,
                blocks,
                page_layout,
                allow_cover_page=page_index == 0,
            )
        for page_layout, section in zip(
            layout.pages,
            document.sections,
            strict=True,
        ):
            if page_layout.page_number_candidate is not None:
                self._set_page_number_start(
                    section,
                    page_layout.page_number_candidate.format.current,
                )

    def _append_visual_pages(
        self,
        document: DocumentType,
        pdf: fitz.Document,
    ) -> None:
        for page_index, page in enumerate(pdf):
            if page_index == 0:
                section = document.sections[0]
            else:
                section = document.add_section(WD_SECTION.NEW_PAGE)
                section_break = document.paragraphs[-1]
                section_break.paragraph_format.space_before = Pt(0)
                section_break.paragraph_format.space_after = Pt(0)
                # Keep the section break itself from consuming visible page
                # space. The image paragraph below deliberately stays on
                # automatic line height so Word cannot clip the drawing.
                section_break.paragraph_format.line_spacing = Pt(1)
            self._configure_visual_section(section, page.rect)
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.keep_together = True
            rendered_page = page.get_pixmap(dpi=150, alpha=False).tobytes("png")
            scale = min(
                (page.rect.width - 2) / page.rect.width,
                (page.rect.height - 8) / page.rect.height,
            )
            paragraph.add_run().add_picture(
                io.BytesIO(rendered_page),
                width=Pt(max(1, page.rect.width * scale)),
                height=Pt(max(1, page.rect.height * scale)),
            )

    @staticmethod
    def _configure_styles(document: DocumentType) -> None:
        update_fields = document.settings._element.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            document.settings._element.append(update_fields)
        update_fields.set(qn("w:val"), "true")
        styles = document.styles
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(11)
        styles["Normal"].paragraph_format.space_after = Pt(2)
        styles["Normal"].paragraph_format.line_spacing = 1.1
        styles["Title"].font.name = "Arial"
        styles["Title"].font.size = Pt(20)
        styles["Title"].font.bold = True
        styles["Title"].paragraph_format.space_before = Pt(0)
        styles["Title"].paragraph_format.space_after = Pt(3)
        styles["Title"].paragraph_format.line_spacing = 1.04
        for style_name, size in (("Heading 1", 15), ("Heading 2", 13)):
            styles[style_name].font.name = "Arial"
            styles[style_name].font.size = Pt(size)
            styles[style_name].font.bold = True
            styles[style_name].paragraph_format.space_before = Pt(4.5)
            styles[style_name].paragraph_format.space_after = Pt(2.5)
            styles[style_name].paragraph_format.line_spacing = 1.06

    @staticmethod
    def _configure_page_size(section: Any, rectangle: fitz.Rect) -> None:
        section.page_width = Pt(rectangle.width)
        section.page_height = Pt(rectangle.height)
        section.orientation = (
            WD_ORIENT.LANDSCAPE
            if rectangle.width > rectangle.height
            else WD_ORIENT.PORTRAIT
        )

    def _configure_editable_section(
        self,
        section: Any,
        page: fitz.Page,
        blocks: Sequence[dict[str, Any]],
        page_layout: PageLayout,
    ) -> None:
        self._configure_page_size(section, page.rect)
        classified_indexes = page_layout.classified_source_indexes
        content_rectangles = [
            fitz.Rect(block["bbox"])
            for block_index, block in enumerate(blocks)
            if block_index not in classified_indexes
            if block.get("bbox") and block.get("type") in {0, 1}
        ]
        if content_rectangles:
            left = min(rectangle.x0 for rectangle in content_rectangles)
            right = page.rect.width - max(
                rectangle.x1 for rectangle in content_rectangles
            )
            top = min(rectangle.y0 for rectangle in content_rectangles)
            bottom = page.rect.height - max(
                rectangle.y1 for rectangle in content_rectangles
            )
        else:
            left = right = top = bottom = 36
        section.left_margin = Pt(self._bounded_margin(left))
        section.right_margin = Pt(self._bounded_margin(right))
        header_blocks = [
            *page_layout.header_candidates,
            *page_layout.header_images,
        ]
        header_bottom = max(
            (block.bbox[3] for block in header_blocks),
            default=0,
        )
        section.top_margin = Pt(
            max(
                self._bounded_margin(top),
                min(top, header_bottom + 8),
            )
        )
        section.bottom_margin = Pt(self._bounded_margin(bottom))
        section.header_distance = Pt(
            min(
                36,
                max(
                    0,
                    min(
                        (block.bbox[1] for block in header_blocks),
                        default=0,
                    ),
                ),
            )
        )
        footer_blocks = [*page_layout.footer_candidates]
        footer_blocks.extend(page_layout.footer_images)
        if (
            page_layout.page_number_candidate is not None
            and page_layout.page_number_candidate.block.bbox[1]
            >= page_layout.height / 2
        ):
            footer_blocks.append(page_layout.page_number_candidate.block)
        section.footer_distance = Pt(
            min(
                36,
                max(
                    0,
                    min(
                        (
                            page_layout.height - block.bbox[3]
                            for block in footer_blocks
                        ),
                        default=0,
                    ),
                ),
            )
        )

    def _configure_header_footer(
        self,
        section: Any,
        page: fitz.Page,
        page_layout: PageLayout,
    ) -> None:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        self._clear_story(section.header)
        self._clear_story(section.footer)

        self._append_story_content(
            section,
            section.header,
            page,
            page_layout.header_candidates,
            page_layout.header_images,
        )
        self._append_story_content(
            section,
            section.footer,
            page,
            page_layout.footer_candidates,
            page_layout.footer_images,
        )

        page_number = page_layout.page_number_candidate
        if page_number is not None:
            story = (
                section.header
                if page_number.block.bbox[3] <= page_layout.height / 2
                else section.footer
            )
            self._append_page_number(story, page_number)

        if not section.header.paragraphs:
            section.header.add_paragraph()
        if not section.footer.paragraphs:
            section.footer.add_paragraph()

    @staticmethod
    def _clear_story(story: Any) -> None:
        for paragraph in list(story.paragraphs):
            paragraph._element.getparent().remove(paragraph._element)

    def _append_story_content(
        self,
        section: Any,
        story: Any,
        page: fitz.Page,
        blocks: Sequence[LayoutBlock],
        images: Sequence[ImageLayoutBlock],
    ) -> None:
        entries: list[tuple[float, float, str, Any]] = [
            (image.bbox[1], image.bbox[0], "image", image)
            for image in images
        ]
        raw_lines = [line for block in blocks for line in block.raw_lines]
        canonicalized = canonicalize_visual_lines(
            raw_lines,
            page.rect.width,
            origin=getattr(self, "_text_origin", "native"),
        )
        self.last_flow_metrics["raw_text_fragments"] += (
            canonicalized.raw_fragment_count
        )
        self.last_flow_metrics["logical_lines"] += len(canonicalized.lines)
        self.last_flow_metrics["same_baseline_fragments_merged"] += (
            canonicalized.same_baseline_fragments_merged
        )
        layouts = analyze_paragraph_layouts(
            canonicalized.lines,
            page.rect,
            page.rect,
            origin=getattr(self, "_text_origin", "native"),
        )
        self.last_flow_metrics["paragraphs"] += len(layouts)
        entries.extend(
            (
                layout.bbox[1],
                layout.bbox[0],
                "paragraph",
                (
                    layout,
                    self._story_layout_alignment(layout, blocks),
                ),
            )
            for layout in layouts
        )
        for _top, _left, entry_type, payload in sorted(
            entries,
            key=lambda entry: entry[:2],
        ):
            if entry_type == "image":
                self._append_story_image(section, story, page, payload)
            else:
                layout, alignment = payload
                self._append_story_paragraph_layout(story, layout, alignment)

    @staticmethod
    def _story_layout_alignment(
        layout: ParagraphLayout,
        blocks: Sequence[LayoutBlock],
    ) -> str:
        source_indexes = {
            source_index
            for line in layout.lines
            for source_index in line.source_block_indexes
        }
        alignments = [
            block.alignment
            for block in blocks
            if block.source_index in source_indexes
        ]
        return Counter(alignments).most_common(1)[0][0] if alignments else layout.alignment

    def _append_story_paragraph_layout(
        self,
        story: Any,
        layout: ParagraphLayout,
        alignment: str,
    ) -> None:
        paragraph = story.add_paragraph()
        paragraph.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }[alignment]
        paragraph.paragraph_format.space_before = Pt(layout.space_before)
        paragraph.paragraph_format.space_after = Pt(layout.space_after)
        paragraph.paragraph_format.line_spacing = layout.line_spacing_ratio
        paragraph.paragraph_format.widow_control = True
        previous_text = ""
        for line_index, line in enumerate(layout.lines):
            if line_index and not previous_text.rstrip().endswith("-"):
                paragraph.add_run(" ")
            for span in line.spans:
                text = str(span.get("text", ""))
                self._append_styled_run(
                    paragraph,
                    span,
                    text,
                    highlighted=False,
                )
                previous_text = text

    def _append_story_image(
        self,
        section: Any,
        story: Any,
        page: fitz.Page,
        image: ImageLayoutBlock,
    ) -> None:
        rectangle = fitz.Rect(image.bbox)
        try:
            rendered_image = self._render_image(
                page,
                image.source_block,
                rectangle,
            )
            available_width = (
                section.page_width - section.left_margin - section.right_margin
            ) / 12700
            scale = min(1.0, available_width / rectangle.width)
            paragraph = story.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            self._apply_story_alignment(paragraph, image.alignment)
            if image.alignment == "left":
                paragraph.paragraph_format.left_indent = Pt(
                    max(0, rectangle.x0 - section.left_margin.pt)
                )
            paragraph.add_run().add_picture(
                io.BytesIO(rendered_image),
                width=Pt(rectangle.width * scale),
                height=Pt(rectangle.height * scale),
            )
        except Exception:
            LOGGER.warning(
                "DOCX header/footer image insertion failed; preserving a marker",
                exc_info=True,
            )
            story.add_paragraph("[Image non convertible]")

    def _append_page_number(
        self,
        story: Any,
        candidate: PageNumberCandidate,
    ) -> None:
        paragraph = story.add_paragraph()
        self._apply_story_alignment(paragraph, candidate.block.alignment)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        page_format = candidate.format
        if page_format.prefix:
            prefix_run = paragraph.add_run(page_format.prefix)
            self._apply_layout_style(prefix_run, candidate.block)
        self._append_word_field(
            paragraph,
            "PAGE",
            str(page_format.current),
            candidate.block,
        )
        if page_format.uses_total:
            separator_run = paragraph.add_run(page_format.separator)
            self._apply_layout_style(separator_run, candidate.block)
            self._append_word_field(
                paragraph,
                "NUMPAGES",
                str(page_format.total),
                candidate.block,
            )
        if page_format.suffix:
            suffix_run = paragraph.add_run(page_format.suffix)
            self._apply_layout_style(suffix_run, candidate.block)

    def _append_word_field(
        self,
        paragraph: Any,
        instruction: str,
        displayed_value: str,
        block: LayoutBlock,
    ) -> None:
        run = paragraph.add_run()
        self._apply_layout_style(run, block)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction_text = OxmlElement("w:instrText")
        instruction_text.set(
            "{http://www.w3.org/XML/1998/namespace}space",
            "preserve",
        )
        instruction_text.text = f" {instruction} "
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        display = OxmlElement("w:t")
        display.text = displayed_value
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend((begin, instruction_text, separate, display, end))

    @staticmethod
    def _apply_story_alignment(paragraph: Any, alignment: str) -> None:
        paragraph.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }[alignment]

    @staticmethod
    def _apply_layout_style(run: Any, block: LayoutBlock) -> None:
        run.bold = block.style.bold
        run.italic = block.style.italic
        run.font.name = block.style.font_family
        run.font.size = Pt(min(40, max(6, block.style.font_size)))
        run.font.color.rgb = RGBColor(
            (block.style.color >> 16) & 0xFF,
            (block.style.color >> 8) & 0xFF,
            block.style.color & 0xFF,
        )

    @staticmethod
    def _set_page_number_start(
        section: Any,
        start_number: int,
    ) -> None:
        section_properties = section._sectPr
        page_number_type = section_properties.find(qn("w:pgNumType"))
        if page_number_type is None:
            page_number_type = OxmlElement("w:pgNumType")
            section_properties.append(page_number_type)
        page_number_type.set(qn("w:start"), str(start_number))

    def _configure_visual_section(
        self,
        section: Any,
        rectangle: fitz.Rect,
    ) -> None:
        self._configure_page_size(section, rectangle)
        section.left_margin = Pt(0)
        section.right_margin = Pt(0)
        section.top_margin = Pt(0)
        section.bottom_margin = Pt(0)
        section.header_distance = Pt(0)
        section.footer_distance = Pt(0)

    @staticmethod
    def _bounded_margin(value: float) -> float:
        return min(54, max(24, value))

    def _append_page_content(
        self,
        document: DocumentType,
        page: fitz.Page,
        blocks: Sequence[dict[str, Any]],
        page_layout: PageLayout,
        *,
        allow_cover_page: bool = False,
    ) -> None:
        classified_indexes = page_layout.classified_source_indexes
        body_source_blocks = [
            {
                **block,
                "_source_index": block_index,
                "_page_index": page_layout.page_index,
                "_story_type": "body",
                "_origin": getattr(self, "_text_origin", "native"),
            }
            for block_index, block in enumerate(blocks)
            if block_index not in classified_indexes
        ]
        cover_analysis = self._analyze_cover_page(
            page,
            body_source_blocks,
            allow_cover_page=allow_cover_page,
        )
        regular_size = self._regular_font_size(body_source_blocks)
        text_content_rectangle = self._text_content_rectangle(
            body_source_blocks,
            page.rect,
        )
        yellow_rectangles, border_rectangles = self._drawing_rectangles(page)
        table_entries, table_rectangles = self._table_entries(page)
        page_has_text = any(
            block.get("type") == 0
            and any(
                span.get("text", "").strip()
                for line in block.get("lines", [])
                for span in line.get("spans", [])
            )
            for block in body_source_blocks
        )
        seen_images: set[tuple[int, tuple[int, int, int, int]]] = set()

        entries: list[tuple[float, float, str, Any]] = list(table_entries)
        raw_text_lines: list[dict[str, Any]] = []
        for block in body_source_blocks:
            source_block_index = int(block["_source_index"])
            block_rectangle = fitz.Rect(block.get("bbox", (0, 0, 0, 0)))
            if block.get("type") == 0 and any(
                rectangle.contains(block_rectangle)
                for rectangle in table_rectangles
            ):
                continue
            if block.get("type") == 0:
                raw_text_lines.extend(
                    {
                        **line,
                        "_source_block_index": source_block_index,
                        "_page_index": page_layout.page_index,
                        "_story_type": "body",
                        "_story_confidence": 1,
                        "_origin": getattr(self, "_text_origin", "native"),
                    }
                    for line in block.get("lines", [])
                    if any(
                        str(span.get("text", "")).strip()
                        for span in line.get("spans", [])
                    )
                )
                continue
            if block.get("type") == 1:
                if (
                    page_has_text
                    and block_rectangle.get_area()
                    >= page.rect.get_area() * 0.65
                ):
                    continue
                image_content = block.get("image")
                image_key = (
                    hash(image_content) if isinstance(image_content, bytes) else 0,
                    tuple(round(value) for value in block_rectangle),
                )
                if image_key in seen_images:
                    continue
                seen_images.add(image_key)
                entries.append(
                    (
                        block_rectangle.y0,
                        block_rectangle.x0,
                        "image",
                        block,
                    )
                )

        canonicalized = canonicalize_visual_lines(
            raw_text_lines,
            page.rect.width,
            origin=getattr(self, "_text_origin", "native"),
        )
        self.last_flow_metrics["raw_text_fragments"] += (
            canonicalized.raw_fragment_count
        )
        self.last_flow_metrics["logical_lines"] += len(canonicalized.lines)
        self.last_flow_metrics["same_baseline_fragments_merged"] += (
            canonicalized.same_baseline_fragments_merged
        )
        paragraph_layouts = analyze_paragraph_layouts(
            canonicalized.lines,
            page.rect,
            text_content_rectangle,
            origin=getattr(self, "_text_origin", "native"),
        )
        self.last_flow_metrics["paragraphs"] += len(paragraph_layouts)
        for layout_index, paragraph_layout in enumerate(paragraph_layouts):
            rectangle = fitz.Rect(paragraph_layout.bbox)
            previous_is_list = (
                layout_index > 0
                and self._is_list_line(paragraph_layouts[layout_index - 1].lines[0])
            )
            next_is_list = (
                layout_index + 1 < len(paragraph_layouts)
                and self._is_list_line(paragraph_layouts[layout_index + 1].lines[0])
            )
            entries.append(
                (
                    rectangle.y0,
                    rectangle.x0,
                    "paragraph",
                    (
                        paragraph_layout,
                        not previous_is_list,
                        not next_is_list,
                    ),
                )
            )

        previous_bottom: float | None = None
        for _top, _left, entry_type, payload in sorted(
            entries,
            key=lambda entry: entry[:2],
        ):
            entry_rectangle = self._entry_rectangle(entry_type, payload)
            paragraph_count = len(document.paragraphs)
            if entry_type == "table":
                self._append_table(document, payload)
            elif entry_type == "image":
                self._append_image_block(document, page, payload)
            elif entry_type == "paragraph":
                paragraph_layout, starts_list, ends_list = payload
                self._append_paragraph_layout(
                    document,
                    page,
                    paragraph_layout,
                    regular_size=regular_size,
                    yellow_rectangles=yellow_rectangles,
                    border_rectangles=border_rectangles,
                    starts_list=starts_list,
                    ends_list=ends_list,
                )
            if (
                cover_analysis.is_cover_page
                and len(document.paragraphs) > paragraph_count
            ):
                previous_source_position = (
                    previous_bottom
                    if previous_bottom is not None
                    else document.sections[-1].top_margin.pt
                )
                self._apply_cover_page_gap(
                    document.paragraphs[paragraph_count],
                    entry_rectangle.y0 - previous_source_position,
                    page.rect.height,
                )
            elif (
                previous_bottom is not None
                and entry_type == "paragraph"
                and len(document.paragraphs) > paragraph_count
            ):
                source_gap = max(0, entry_rectangle.y0 - previous_bottom)
                space_before = normalize_paragraph_gap(
                    source_gap,
                    regular_size,
                )
                if space_before:
                    first_paragraph = document.paragraphs[paragraph_count]
                    existing = first_paragraph.paragraph_format.space_before
                    existing_points = existing.pt if existing is not None else 0
                    first_paragraph.paragraph_format.space_before = Pt(
                        max(existing_points, space_before)
                    )
            previous_bottom = (
                entry_rectangle.y1
                if previous_bottom is None
                else max(previous_bottom, entry_rectangle.y1)
            )

    @classmethod
    def _analyze_cover_page(
        cls,
        page: fitz.Page,
        blocks: list[dict[str, Any]],
        *,
        allow_cover_page: bool,
    ) -> CoverPageAnalysis:
        text_blocks = [
            block
            for block in blocks
            if block.get("type") == 0
            and cls._block_text(block).strip()
        ]
        text = " ".join(cls._block_text(block) for block in text_blocks)
        word_count = len(re.findall(r"\w+", text))
        text_rectangles = [
            fitz.Rect(block["bbox"])
            for block in text_blocks
            if block.get("bbox")
        ]
        content_rectangles = [
            fitz.Rect(block["bbox"])
            for block in blocks
            if block.get("bbox")
            and block.get("type") in {0, 1}
            and fitz.Rect(block["bbox"]).get_area()
            < page.rect.get_area() * 0.65
        ]
        ordered_rectangles = sorted(
            content_rectangles,
            key=lambda rectangle: (rectangle.y0, rectangle.x0),
        )
        vertical_gaps = [
            max(0, current.y0 - previous.y1)
            for previous, current in zip(
                ordered_rectangles,
                ordered_rectangles[1:],
                strict=False,
            )
        ]
        largest_vertical_gap = max(vertical_gaps, default=0)
        occupied_area = sum(
            min(rectangle.get_area(), page.rect.get_area())
            for rectangle in content_rectangles
        )
        whitespace_ratio = max(
            0,
            1 - min(1, occupied_area / max(1, page.rect.get_area())),
        )
        span_sizes = [
            float(span.get("size", 0))
            for block in text_blocks
            for line in block.get("lines", [])
            for span in line.get("spans", [])
            if str(span.get("text", "")).strip()
        ]
        largest_font_size = max(span_sizes, default=0)
        regular_font_size = statistics.median(span_sizes) if span_sizes else 11
        title_rectangle = fitz.Rect()
        title_size = 0.0
        for block in text_blocks:
            block_size = max(
                (
                    float(span.get("size", 0))
                    for line in block.get("lines", [])
                    for span in line.get("spans", [])
                    if str(span.get("text", "")).strip()
                ),
                default=0,
            )
            if block_size > title_size:
                title_size = block_size
                title_rectangle = fitz.Rect(block["bbox"])
        title_position_ratio = (
            title_rectangle.y0 / page.rect.height
            if not title_rectangle.is_empty
            else 0
        )
        has_lower_page_text = any(
            rectangle.y0 >= page.rect.height * 0.65
            for rectangle in text_rectangles
        )
        has_long_paragraph = any(
            len(re.findall(r"\w+", cls._block_text(block))) > 45
            for block in text_blocks
        )
        has_list = any(
            cls._is_list_line(line)
            for block in text_blocks
            for line in block.get("lines", [])
        )
        is_cover_page = (
            allow_cover_page
            and 3 <= len(text_blocks) <= 8
            and word_count <= 150
            and not has_long_paragraph
            and not has_list
            and largest_font_size >= max(18, regular_font_size * 1.35)
            and largest_vertical_gap >= page.rect.height * 0.12
            and has_lower_page_text
            and whitespace_ratio >= 0.65
        )
        return CoverPageAnalysis(
            is_cover_page=is_cover_page,
            text_block_count=len(text_blocks),
            word_count=word_count,
            whitespace_ratio=round(whitespace_ratio, 4),
            largest_font_size=largest_font_size,
            largest_vertical_gap=largest_vertical_gap,
            title_position_ratio=round(title_position_ratio, 4),
            has_lower_page_text=has_lower_page_text,
        )

    @staticmethod
    def _block_text(block: dict[str, Any]) -> str:
        return " ".join(
            str(span.get("text", "")).strip()
            for line in block.get("lines", [])
            for span in line.get("spans", [])
            if str(span.get("text", "")).strip()
        )

    @staticmethod
    def _entry_rectangle(entry_type: str, payload: Any) -> fitz.Rect:
        if entry_type == "table":
            return fitz.Rect(payload.bbox)
        if entry_type == "paragraph":
            return fitz.Rect(payload[0].bbox)
        return fitz.Rect(payload.get("bbox", (0, 0, 0, 0)))

    @staticmethod
    def _apply_cover_page_gap(
        paragraph: Any,
        source_gap: float,
        page_height: float,
    ) -> None:
        if source_gap < page_height * 0.035:
            return
        preserved_gap = min(
            source_gap * 0.94,
            page_height * 0.34,
        )
        existing = paragraph.paragraph_format.space_before
        existing_points = existing.pt if existing is not None else 0
        paragraph.paragraph_format.space_before = Pt(
            max(existing_points, preserved_gap),
        )

    @staticmethod
    def _text_content_rectangle(
        blocks: Sequence[dict[str, Any]],
        page_rectangle: fitz.Rect,
    ) -> fitz.Rect:
        rectangles = [
            fitz.Rect(block["bbox"])
            for block in blocks
            if block.get("type") == 0 and block.get("bbox")
        ]
        if not rectangles:
            return fitz.Rect(page_rectangle)
        content_rectangle = fitz.Rect(rectangles[0])
        for rectangle in rectangles[1:]:
            content_rectangle |= rectangle
        return content_rectangle

    @staticmethod
    def _first_text_rectangle(block: dict[str, Any]) -> fitz.Rect:
        for line in block.get("lines", []):
            if any(
                span.get("text", "").strip()
                for span in line.get("spans", [])
            ):
                return fitz.Rect(line["bbox"])
        return fitz.Rect(block.get("bbox", (0, 0, 0, 0)))

    @staticmethod
    def _regular_font_size(blocks: list[dict[str, Any]]) -> float:
        font_sizes = [
            float(span.get("size", 0))
            for block in blocks
            if block.get("type") == 0
            for line in block.get("lines", [])
            for span in line.get("spans", [])
            if span.get("text", "").strip()
        ]
        return statistics.median(font_sizes) if font_sizes else 11

    @staticmethod
    def _drawing_rectangles(
        page: fitz.Page,
    ) -> tuple[list[fitz.Rect], list[fitz.Rect]]:
        yellow_rectangles: list[fitz.Rect] = []
        border_rectangles: list[fitz.Rect] = []
        dark_segments: list[fitz.Rect] = []
        for drawing in page.get_drawings():
            drawing_rectangle = fitz.Rect(drawing.get("rect", (0, 0, 0, 0)))
            fill = drawing.get("fill")
            if (
                not drawing_rectangle.is_empty
                and fill
                and fill[0] >= 0.75
                and fill[1] >= 0.7
                and fill[2] <= 0.55
            ):
                # Highlights exported by office suites are often paths made of
                # curves and lines, not a PDF ``re`` rectangle.
                yellow_rectangles.append(drawing_rectangle)
                continue
            for item in drawing.get("items", []):
                if not item or item[0] != "re":
                    continue
                rectangle = fitz.Rect(item[1])
                if (
                    drawing.get("color") is not None
                    and rectangle.width < page.rect.width * 0.95
                    and rectangle.height < page.rect.height * 0.8
                ):
                    border_rectangles.append(rectangle)
                elif (
                    fill
                    and max(fill) <= 0.2
                    and min(rectangle.width, rectangle.height) <= 2
                ):
                    # Some producers flatten a border into many thin, filled
                    # black rectangles. Reassemble those segments below.
                    dark_segments.append(rectangle)
        border_rectangles.extend(
            PdfToDocxConverter._border_rectangles_from_segments(
                dark_segments,
                page.rect,
            )
        )
        return yellow_rectangles, border_rectangles

    @staticmethod
    def _border_rectangles_from_segments(
        segments: list[fitz.Rect],
        page_rectangle: fitz.Rect,
    ) -> list[fitz.Rect]:
        remaining = list(segments)
        candidates: list[fitz.Rect] = []
        while remaining:
            component = [remaining.pop()]
            changed = True
            while changed:
                changed = False
                for rectangle in list(remaining):
                    if any(
                        PdfToDocxConverter._rectangles_touch(
                            rectangle,
                            member,
                        )
                        for member in component
                    ):
                        component.append(rectangle)
                        remaining.remove(rectangle)
                        changed = True
            bounds = fitz.Rect(component[0])
            for rectangle in component[1:]:
                bounds |= rectangle
            horizontal = [
                rectangle
                for rectangle in component
                if rectangle.width >= bounds.width * 0.7
                and rectangle.height <= 2
            ]
            vertical = [
                rectangle
                for rectangle in component
                if rectangle.height > rectangle.width
                and rectangle.width <= 2
            ]
            has_left = any(
                abs(rectangle.x0 - bounds.x0) <= 2
                for rectangle in vertical
            )
            has_right = any(
                abs(rectangle.x1 - bounds.x1) <= 2
                for rectangle in vertical
            )
            if (
                bounds.width >= 50
                and bounds.height >= 10
                and bounds.width < page_rectangle.width * 0.95
                and bounds.height < page_rectangle.height * 0.8
                and len(horizontal) >= 2
                and has_left
                and has_right
            ):
                candidates.append(bounds)
        return candidates

    @staticmethod
    def _rectangles_touch(
        first: fitz.Rect,
        second: fitz.Rect,
        tolerance: float = 1,
    ) -> bool:
        expanded = fitz.Rect(
            first.x0 - tolerance,
            first.y0 - tolerance,
            first.x1 + tolerance,
            first.y1 + tolerance,
        )
        return expanded.intersects(second)

    @staticmethod
    def _table_entries(
        page: fitz.Page,
    ) -> tuple[list[tuple[float, float, str, Any]], list[fitz.Rect]]:
        try:
            tables = page.find_tables().tables
        except (AttributeError, RuntimeError, ValueError):
            return [], []
        entries = [
            (table.bbox[1], table.bbox[0], "table", table)
            for table in tables
        ]
        return entries, [fitz.Rect(table.bbox) for table in tables]

    @staticmethod
    def _append_table(document: DocumentType, detected_table: Any) -> None:
        cells = detected_table.extract()
        if not cells:
            return
        column_count = max((len(row) for row in cells), default=0)
        if column_count == 0:
            return
        table = document.add_table(rows=len(cells), cols=column_count)
        table.style = "Table Grid"
        for row_index, row in enumerate(cells):
            for column_index, value in enumerate(row):
                table.cell(row_index, column_index).text = value or ""

    def _append_image_block(
        self,
        document: DocumentType,
        page: fitz.Page,
        block: dict[str, Any],
    ) -> None:
        rectangle = fitz.Rect(block["bbox"])
        if rectangle.is_empty or rectangle.width <= 0 or rectangle.height <= 0:
            return
        try:
            rendered_image = self._render_image(page, block, rectangle)
            section = document.sections[-1]
            available_width = (
                section.page_width - section.left_margin - section.right_margin
            ) / 12700
            scale = min(1.0, available_width / rectangle.width)
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.keep_with_next = True
            if self._is_centered(rectangle, page.rect):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                content_left = section.left_margin.pt
                paragraph.paragraph_format.left_indent = Pt(
                    max(0, rectangle.x0 - content_left)
                )
            paragraph.add_run().add_picture(
                io.BytesIO(rendered_image),
                width=Pt(rectangle.width * scale),
                height=Pt(rectangle.height * scale),
            )
        except Exception:
            LOGGER.warning(
                "DOCX image insertion failed; conversion continues with a placeholder",
                exc_info=True,
            )
            document.add_paragraph("[Image non convertible]")

    @classmethod
    def _render_image(
        cls,
        page: fitz.Page,
        block: dict[str, Any],
        rectangle: fitz.Rect,
    ) -> bytes:
        image = block.get("image")
        mask = block.get("mask")
        if isinstance(image, bytes):
            try:
                image_pixmap = cls._to_rgb(fitz.Pixmap(image))
                if isinstance(mask, bytes):
                    # PyMuPDF rejects Pixmap(color, mask) when ``color``
                    # already owns an alpha channel. PDF image blocks can
                    # expose both a decoded alpha channel and the original
                    # soft mask, notably in Firefox regression documents.
                    color_pixmap = cls._drop_alpha(image_pixmap)
                    mask_pixmap = cls._normalize_mask(fitz.Pixmap(mask))
                    image_pixmap = fitz.Pixmap(color_pixmap, mask_pixmap)
                return image_pixmap.tobytes("png")
            except Exception:
                LOGGER.warning(
                    "DOCX image extraction failed; rasterizing its page rectangle",
                    exc_info=True,
                )
        return cls._render_page_clip_on_white(page, rectangle)

    @staticmethod
    def _drop_alpha(pixmap: fitz.Pixmap) -> fitz.Pixmap:
        return fitz.Pixmap(pixmap, 0) if pixmap.alpha else pixmap

    @staticmethod
    def _to_rgb(pixmap: fitz.Pixmap) -> fitz.Pixmap:
        if pixmap.colorspace is None:
            raise ValueError("Une image couleur doit avoir un espace colorimétrique.")
        if pixmap.colorspace.n != 3:
            return fitz.Pixmap(fitz.csRGB, pixmap)
        return pixmap

    @staticmethod
    def _normalize_mask(pixmap: fitz.Pixmap) -> fitz.Pixmap:
        if pixmap.colorspace is not None and pixmap.colorspace.n != 1:
            pixmap = fitz.Pixmap(fitz.csGRAY, pixmap)
        return fitz.Pixmap(pixmap, 0) if pixmap.alpha else pixmap

    @staticmethod
    def _render_page_clip_on_white(
        page: fitz.Page,
        rectangle: fitz.Rect,
    ) -> bytes:
        # alpha=False composites transparency on white instead of black.
        return page.get_pixmap(
            matrix=fitz.Matrix(2, 2),
            clip=rectangle,
            alpha=False,
        ).tobytes("png")

    def _append_text_block(
        self,
        document: DocumentType,
        page: fitz.Page,
        block: dict[str, Any],
        *,
        content_rectangle: fitz.Rect,
        regular_size: float,
        yellow_rectangles: list[fitz.Rect],
        border_rectangles: list[fitz.Rect],
        preserve_vertical_spacing: bool = False,
    ) -> None:
        lines = [
            line
            for line in block.get("lines", [])
            if any(span.get("text", "").strip() for span in line.get("spans", []))
        ]
        if not lines:
            return
        paragraph_layouts = analyze_paragraph_layouts(
            lines,
            page.rect,
            content_rectangle,
            origin=getattr(self, "_text_origin", "native"),
        )
        previous_bottom: float | None = None
        for group_index, paragraph_layout in enumerate(paragraph_layouts):
            group_rectangle = fitz.Rect(paragraph_layout.bbox)
            paragraph_count = len(document.paragraphs)
            self._append_paragraph_layout(
                document,
                page,
                paragraph_layout,
                regular_size=regular_size,
                yellow_rectangles=yellow_rectangles,
                border_rectangles=border_rectangles,
                starts_list=(
                    group_index == 0
                    or not self._is_list_line(
                        paragraph_layouts[group_index - 1].lines[0]
                    )
                ),
                ends_list=(
                    group_index == len(paragraph_layouts) - 1
                    or not self._is_list_line(
                        paragraph_layouts[group_index + 1].lines[0]
                    )
                ),
            )
            if (
                preserve_vertical_spacing
                and previous_bottom is not None
                and len(document.paragraphs) > paragraph_count
            ):
                self._apply_cover_page_gap(
                    document.paragraphs[paragraph_count],
                    group_rectangle.y0 - previous_bottom,
                    page.rect.height,
                )
            previous_bottom = group_rectangle.y1

    def _append_paragraph_layout(
        self,
        document: DocumentType,
        page: fitz.Page,
        layout: ParagraphLayout,
        *,
        regular_size: float,
        yellow_rectangles: list[fitz.Rect],
        border_rectangles: list[fitz.Rect],
        starts_list: bool,
        ends_list: bool,
    ) -> None:
        line_group = list(layout.lines)
        if self._is_list_line(line_group[0]):
            self._append_list_lines(
                document,
                line_group,
                yellow_rectangles=yellow_rectangles,
                starts_list=starts_list,
                ends_list=ends_list,
            )
            return
        self._append_text_paragraph(
            document,
            page,
            layout,
            regular_size=regular_size,
            yellow_rectangles=yellow_rectangles,
            border_rectangles=border_rectangles,
        )

    def _append_text_paragraph(
        self,
        document: DocumentType,
        page: fitz.Page,
        layout: ParagraphLayout,
        *,
        regular_size: float,
        yellow_rectangles: list[fitz.Rect],
        border_rectangles: list[fitz.Rect],
    ) -> None:
        lines = list(layout.lines)
        block_rectangle = fitz.Rect(layout.bbox)
        largest_size = max(
            float(span.get("size", 0))
            for line in lines
            for span in line.get("spans", [])
        )
        text_length = sum(
            len(span.get("text", ""))
            for line in lines
            for span in line.get("spans", [])
        )
        bold_ratio = self._bold_text_ratio(lines)
        centered = layout.alignment == "center"
        style = self._paragraph_style(
            largest_size,
            regular_size,
            centered=centered,
            text_length=text_length,
            bold_ratio=bold_ratio,
        )
        paragraph = document.add_paragraph(style=style)
        paragraph.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }[layout.alignment]
        paragraph.paragraph_format.keep_together = style != "Normal"
        paragraph.paragraph_format.keep_with_next = style != "Normal"
        paragraph.paragraph_format.widow_control = True
        if layout.left_indent:
            paragraph.paragraph_format.left_indent = Pt(layout.left_indent)
        if layout.right_indent:
            paragraph.paragraph_format.right_indent = Pt(layout.right_indent)
        if layout.first_line_indent:
            paragraph.paragraph_format.first_line_indent = Pt(
                layout.first_line_indent
            )
        if style == "Title":
            paragraph.paragraph_format.space_before = Pt(layout.space_before)
            paragraph.paragraph_format.space_after = Pt(
                max(3, layout.space_after)
            )
            paragraph.paragraph_format.line_spacing = 1.05
        elif style.startswith("Heading"):
            paragraph.paragraph_format.space_before = Pt(
                max(3, layout.space_before)
            )
            paragraph.paragraph_format.space_after = Pt(
                max(2.5, layout.space_after)
            )
            paragraph.paragraph_format.line_spacing = max(
                1.05,
                layout.line_spacing_ratio,
            )
        else:
            paragraph.paragraph_format.space_before = Pt(layout.space_before)
            paragraph.paragraph_format.space_after = Pt(layout.space_after)
            paragraph.paragraph_format.line_spacing = (
                1.15
                if len(lines) < 2 and text_length >= 120
                else layout.line_spacing_ratio
            )

        previous_text = ""
        for line_index, line in enumerate(lines):
            spans = line.get("spans", [])
            if line_index:
                separator = "" if previous_text.rstrip().endswith("-") else " "
                if separator:
                    paragraph.add_run(separator)
            for span in spans:
                text = str(span.get("text", ""))
                self._append_styled_run(
                    paragraph,
                    span,
                    text,
                    highlighted=self._is_highlighted(
                        fitz.Rect(span.get("bbox", block_rectangle)),
                        yellow_rectangles,
                    ),
                )
                previous_text = text

        if any(
            self._rectangle_frames_block(rectangle, block_rectangle)
            for rectangle in border_rectangles
        ):
            self._add_paragraph_border(paragraph)

    @staticmethod
    def _bold_text_ratio(lines: list[LogicalLine]) -> float:
        bold_characters = 0
        total_characters = 0
        for line in lines:
            for span in line.get("spans", []):
                text = str(span.get("text", "")).strip()
                if not text:
                    continue
                character_count = len(text)
                total_characters += character_count
                flags = int(span.get("flags", 0))
                font_name = str(span.get("font", "")).lower()
                if bool(flags & 16) or "bold" in font_name:
                    bold_characters += character_count
        return (
            bold_characters / total_characters
            if total_characters
            else 0
        )

    @staticmethod
    def _is_list_line(line: LogicalLine | dict[str, Any]) -> bool:
        return is_list_line(line)

    def _append_list_lines(
        self,
        document: DocumentType,
        lines: list[LogicalLine],
        *,
        yellow_rectangles: list[fitz.Rect],
        starts_list: bool,
        ends_list: bool,
    ) -> None:
        line = lines[0]
        spans = line.get("spans", [])
        text = "".join(str(span.get("text", "")) for span in spans)
        marker = LIST_MARKER.match(text)
        if marker is None:
            return
        item_text = text[marker.end() :].strip()
        continuation_text = " ".join(
            "".join(
                str(span.get("text", ""))
                for span in continuation.get("spans", [])
            ).strip()
            for continuation in lines[1:]
        ).strip()
        if not item_text and not continuation_text:
            return
        numbered = marker.group("marker")[0].isdigit()
        paragraph = document.add_paragraph(
            style="List Number" if numbered else "List Bullet"
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.first_line_indent = Pt(-12)
        paragraph.paragraph_format.space_before = Pt(2.5 if starts_list else 0)
        paragraph.paragraph_format.space_after = Pt(4 if ends_list else 0.75)
        paragraph.paragraph_format.line_spacing = 1.08
        paragraph.paragraph_format.keep_together = False
        paragraph.paragraph_format.widow_control = True
        remaining_prefix = marker.end()
        consumed = 0
        for span in spans:
            span_text = str(span.get("text", ""))
            start = max(0, remaining_prefix - consumed)
            rendered_text = span_text[start:]
            consumed += len(span_text)
            if not rendered_text:
                continue
            self._append_styled_run(
                paragraph,
                span,
                rendered_text,
                highlighted=self._is_highlighted(
                    fitz.Rect(span.get("bbox", (0, 0, 0, 0))),
                    yellow_rectangles,
                ),
            )
        for continuation in lines[1:]:
            paragraph.add_run(" ")
            for span in continuation.get("spans", []):
                self._append_styled_run(
                    paragraph,
                    span,
                    str(span.get("text", "")),
                    highlighted=self._is_highlighted(
                        fitz.Rect(span.get("bbox", (0, 0, 0, 0))),
                        yellow_rectangles,
                    ),
                )

    @staticmethod
    def _paragraph_style(
        largest_size: float,
        regular_size: float,
        *,
        centered: bool,
        text_length: int,
        bold_ratio: float,
    ) -> str:
        if (
            centered
            and text_length < 100
            and largest_size >= regular_size * 1.25
        ):
            return "Title"
        if text_length < 160 and largest_size >= max(18, regular_size * 1.55):
            return "Title" if centered else "Heading 1"
        if text_length < 180 and largest_size >= max(13.5, regular_size * 1.22):
            return "Heading 2"
        if (
            not centered
            and text_length <= 110
            and bold_ratio >= 0.8
        ):
            return "Heading 2"
        return "Normal"

    @staticmethod
    def _append_styled_run(
        paragraph: Any,
        span: dict[str, Any],
        text: str,
        *,
        highlighted: bool,
    ) -> None:
        run = paragraph.add_run(text)
        flags = int(span.get("flags", 0))
        font_name = str(span.get("font", "Arial")).split("+")[-1]
        run.bold = bool(flags & 16) or "bold" in font_name.lower()
        run.italic = bool(flags & 2) or any(
            marker in font_name.lower() for marker in ("italic", "oblique")
        )
        run.font.name = font_name
        run.font.size = Pt(min(40, max(6, float(span.get("size", 11)))))
        color = int(span.get("color", 0))
        run.font.color.rgb = RGBColor(
            (color >> 16) & 0xFF,
            (color >> 8) & 0xFF,
            color & 0xFF,
        )
        if highlighted:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    @staticmethod
    def _is_centered(rectangle: fitz.Rect, page_rectangle: fitz.Rect) -> bool:
        return (
            abs(
                (rectangle.x0 + rectangle.x1) / 2
                - (page_rectangle.x0 + page_rectangle.x1) / 2
            )
            <= 22
            and rectangle.width <= page_rectangle.width * 0.9
        )

    @staticmethod
    def _is_highlighted(
        span_rectangle: fitz.Rect,
        yellow_rectangles: list[fitz.Rect],
    ) -> bool:
        return any(
            rectangle.intersects(span_rectangle)
            and (rectangle & span_rectangle).get_area()
            >= span_rectangle.get_area() * 0.35
            for rectangle in yellow_rectangles
            if span_rectangle.get_area() > 0
        )

    @staticmethod
    def _rectangle_frames_block(
        rectangle: fitz.Rect,
        block_rectangle: fitz.Rect,
    ) -> bool:
        tolerance = 2
        return (
            rectangle.x0 <= block_rectangle.x0 + tolerance
            and rectangle.y0 <= block_rectangle.y0 + tolerance
            and rectangle.x1 >= block_rectangle.x1 - tolerance
            and rectangle.y1 >= block_rectangle.y1 - tolerance
            and rectangle.width <= block_rectangle.width + 80
            and rectangle.height <= block_rectangle.height + 100
        )

    @staticmethod
    def _add_paragraph_border(paragraph: Any) -> None:
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.12
        paragraph_properties = paragraph._p.get_or_add_pPr()
        borders = paragraph_properties.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            paragraph_properties.append(borders)
        for edge_name in ("top", "left", "bottom", "right"):
            edge = OxmlElement(f"w:{edge_name}")
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), "8")
            edge.set(qn("w:space"), "8")
            edge.set(qn("w:color"), "333333")
            borders.append(edge)
